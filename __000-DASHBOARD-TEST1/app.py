import os
from typing import Any
import sys
import yaml  # type: ignore[import-not-found]
import subprocess
import glob
import streamlit as st  # type: ignore[import-not-found]
import streamlit.components.v1 as components  # type: ignore[import-not-found]
import tempfile
import mimetypes
import time
import base64
import shlex

# -----------------------------------------------------------------------------
# Helper: Get Python Path (prefers virtual environment)
# -----------------------------------------------------------------------------
def get_python_cmd():
    # Use the venv Python from the same directory as this app.py file.
    # sys.executable may point to a DIFFERENT venv (e.g. V2 symlink) that
    # lacks required packages like google-api-python-client.
    import sys
    app_dir = os.path.dirname(os.path.abspath(__file__))
    venv_python = os.path.join(app_dir, ".venv", "bin", "python")
    if os.path.exists(venv_python):
        return venv_python
    venv_python3 = os.path.join(app_dir, ".venv", "bin", "python3")
    if os.path.exists(venv_python3):
        return venv_python3
    # Fallback to the running interpreter
    return sys.executable

# -----------------------------------------------------------------------------
# Helper: Load External CSS
# -----------------------------------------------------------------------------
def load_css():
    css_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "style.css")
    if os.path.exists(css_path):
        with open(css_path, "r") as f:
            # Add a timestamp comment to force Streamlit to see this as new content
            css_content = f"/* {time.time()} */\n" + f.read()
            st.markdown(f"<style>{css_content}</style>", unsafe_allow_html=True)

def render_speed_controls(skill_id=None, stats_text="", clip_name=""):
    skill_val = f"'{skill_id}'" if skill_id else "null"
    import html as _html
    safe_clip = _html.escape(clip_name) if clip_name else ""
    stats_html = ""
    if stats_text:
        safe_stats = _html.escape(stats_text)
        stats_html = f"""<span style="
            margin-left: 14px;
            background-color: #000;
            color: #fff;
            font-size: 15px;
            font-weight: 600;
            font-family: sans-serif;
            padding: 3px 12px;
            border-radius: 5px;
            letter-spacing: 0.02em;
            white-space: nowrap;
        ">{safe_stats}</span>"""
    components.html(
        f"""
        <style>
            body {{ margin: 0; padding: 0; background: transparent; }}
            .speed-bar {{
                display: flex;
                flex-direction: row;
                gap: 8px;
                align-items: center;
                justify-content: flex-start;
                padding: 4px 0;
            }}
            .speed-btn {{
                background: #2b2b36;
                color: #fafafa;
                border: 1px solid #454555;
                border-radius: 4px;
                padding: 4px 10px;
                font-family: sans-serif;
                font-size: 13px;
                cursor: pointer;
                transition: all 0.2s;
            }}
            .speed-btn:hover {{
                background: #3b3b4a;
                border-color: #ffd166;
            }}
            .follow-toggle {{
                background: #2b2b36;
                color: #fafafa;
                border: 1px solid #454555;
                border-radius: 4px;
                padding: 4px 12px;
                font-family: sans-serif;
                font-size: 13px;
                cursor: pointer;
                transition: all 0.2s;
                margin-left: 12px;
                white-space: nowrap;
            }}
            .follow-toggle:hover {{
                border-color: #88e0e4;
            }}
            .follow-toggle.active {{
                background: #1a6b6e;
                color: #a8ffdb;
                border-color: #88e0e4;
                font-weight: bold;
            }}
        </style>
        <div class="speed-bar">
            <span style="color: #bfbfbf; font-family: sans-serif; font-size: 13px; margin-right: 4px;">Speed:</span>
            <button class="speed-btn" onclick="setSpeed(0.5)">0.5x</button>
            <button class="speed-btn" onclick="setSpeed(1.0)">1x</button>
            <button class="speed-btn" onclick="setSpeed(1.25)">1.25x</button>
            <button class="speed-btn" onclick="setSpeed(1.5)">1.5x</button>
            <button class="speed-btn" onclick="setSpeed(2.0)">2x</button>
            <button class="speed-btn" onclick="setSpeed(3.0)">3x</button>
            <button class="speed-btn" onclick="setSpeed(4.0)">4x</button>
            {stats_html}
            <button class="follow-toggle" id="follow-toggle-btn" onclick="toggleFollow()">📖 Follow Along</button>
        </div>
        <script>
            const currentSkillId = {skill_val};
            const clipName = '{safe_clip}';
            
            if (currentSkillId) {{
                // sessionStorage only persists during the tab session (lost on close)
                // whereas localStorage persists across browser sessions.
                const sessionSkillId = sessionStorage.getItem('sessionSkillId');
                
                if (sessionSkillId !== currentSkillId) {{
                    // This is either a NEW tab session or a SKILL SWITCH.
                    // Reset to 1x per user preference.
                    localStorage.setItem('audioPlaybackSpeed', '1.0');
                    sessionStorage.setItem('sessionSkillId', currentSkillId);
                    localStorage.setItem('lastSkillId', currentSkillId);
                }}
            }}

            function setSpeed(rate) {{
                localStorage.setItem('audioPlaybackSpeed', rate);
                const audios = window.parent.document.querySelectorAll('audio');
                audios.forEach(a => {{
                    a.playbackRate = rate;
                }});
                updateButtons(rate);
            }}

            function updateButtons(rate) {{
                const btns = document.querySelectorAll('.speed-btn');
                btns.forEach(btn => {{
                    const btnRate = parseFloat(btn.innerText);
                    if (btnRate === rate) {{
                        btn.style.background = '#ffd166';
                        btn.style.color = '#000';
                        btn.style.borderColor = '#ffd166';
                        btn.style.fontWeight = 'bold';
                    }} else {{
                        btn.style.background = '#2b2b36';
                        btn.style.color = '#fafafa';
                        btn.style.borderColor = '#454555';
                        btn.style.fontWeight = 'normal';
                    }}
                }});
            }}

            const applyStoredSpeed = () => {{
                const savedSpeed = parseFloat(localStorage.getItem('audioPlaybackSpeed') || '1.0');
                const volKey = clipName ? ('ag_vol_' + clipName) : 'ag_audio_vol';
                const savedVol = parseFloat(localStorage.getItem(volKey));
                const vol = isNaN(savedVol) ? 0.45 : savedVol;
                const audios = window.parent.document.querySelectorAll('audio');
                audios.forEach(a => {{
                    if (Math.abs(a.playbackRate - savedSpeed) > 0.01) {{
                        a.playbackRate = savedSpeed;
                    }}
                    // Apply saved volume and attach listener (once per element)
                    if (!a._agVolBound) {{
                        a.volume = vol;
                        a._agVolBound = true;
                        a.addEventListener('volumechange', function() {{
                            localStorage.setItem(volKey, String(a.volume));
                        }});
                    }}
                }});
                updateButtons(savedSpeed);
            }};

            // Follow Along toggle
            function toggleFollow() {{
                const btn = document.getElementById('follow-toggle-btn');
                const current = localStorage.getItem('agFollowAlongEnabled');
                const newState = (current === 'false') ? 'true' : 'false';
                localStorage.setItem('agFollowAlongEnabled', newState);
                updateFollowBtn(newState === 'true');
            }}
            function updateFollowBtn(isOn) {{
                const btn = document.getElementById('follow-toggle-btn');
                if (!btn) return;
                if (isOn) {{
                    btn.classList.add('active');
                    btn.textContent = '📖 Follow Along: ON';
                }} else {{
                    btn.classList.remove('active');
                    btn.textContent = '📖 Follow Along: OFF';
                }}
            }}
            // Default to ON if not set
            if (localStorage.getItem('agFollowAlongEnabled') === null) {{
                localStorage.setItem('agFollowAlongEnabled', 'true');
            }}
            updateFollowBtn(localStorage.getItem('agFollowAlongEnabled') !== 'false');

            window.addEventListener('load', applyStoredSpeed);
            applyStoredSpeed();
            setInterval(applyStoredSpeed, 500);
        </script>
        """,
        height=40
    )

def render_word_tracker(alignment_data, clip_name="", estimated=False):
    """Renders a synchronized word-tracking text display that highlights words in sync with audio playback.
    
    If estimated=True, evenly distributes words across audio.duration (for Speech2Text parity).
    """
    import json as _json_wt
    import html as _html_wt
    
    words = alignment_data.get("words", [])
    starts = alignment_data.get("start_times", [])
    ends = alignment_data.get("end_times", [])
    
    if not words:
        return
    
    # Build word spans with data attributes for timing
    word_spans = []
    for i, word in enumerate(words):
        safe_word = _html_wt.escape(word)
        if estimated:
            # No timing data — JS will compute from audio.duration
            word_spans.append(f'<span class="wt-word" data-idx="{i}">{safe_word}</span>')
        else:
            s = starts[i] if i < len(starts) else 0
            e = ends[i] if i < len(ends) else 0
            word_spans.append(f'<span class="wt-word" data-idx="{i}" data-start="{s}" data-end="{e}">{safe_word}</span>')
    
    words_html = " ".join(word_spans)
    
    # Calculate appropriate height for the component iframe
    # At 31px font, ~5 words per line centered, 62px per line (31*2.0), plus 76px padding (38*2)
    words_per_line = max(1, 6)
    estimated_lines = max(1, (len(words) + words_per_line - 1) // words_per_line)
    container_height = estimated_lines * 45 + 64
    
    components.html(f"""
        <style>
            body {{ margin: 0; padding: 0; background: transparent; }}
            .wt-container {{
                font-family: inherit;
                font-size: 28px;
                line-height: 1.5;
                color: #ffffff;
                padding: 19px 25px;
                background-color: #000000;
                border-radius: 10px;
                border: 2px solid #ffffff;
                text-align: center;
                margin-bottom: 0;
                word-spacing: -2px;
            }}
            .wt-container::-webkit-scrollbar {{ width: 6px; }}
            .wt-container::-webkit-scrollbar-track {{ background: transparent; }}
            .wt-container::-webkit-scrollbar-thumb {{ background: rgba(255, 255, 255, 0.3); border-radius: 3px; }}
            .wt-word {{
                display: inline;
                padding: 3px 5px;
                border-radius: 4px;
                transition: all 0.18s ease;
                cursor: default;
            }}
            .wt-word.active {{
                background: rgba(212, 249, 250, 0.4);
                color: #d4f9fa;
                text-shadow: -1px 0 0 currentColor, 1px 0 0 currentColor;
                display: inline-block;
                transform: scale(1.20);
            }}
            .wt-word.past {{
                color: #777777;
            }}
        </style>
        <div class="wt-container" id="wt-scroll-container">
            {words_html}
        </div>
        <script>
            const container = document.getElementById('wt-scroll-container');
            const wordEls = container.querySelectorAll('.wt-word');
            let lastActiveIdx = -1;
            const isEstimated = {'true' if estimated else 'false'};
            let estimatedTimingsSet = false;
            
            function setEstimatedTimings(duration) {{
                if (estimatedTimingsSet || !isEstimated || !duration) return;
                estimatedTimingsSet = true;
                const totalWords = wordEls.length;
                const timePerWord = duration / totalWords;
                wordEls.forEach((el, idx) => {{
                    el.dataset.start = (idx * timePerWord).toFixed(4);
                    el.dataset.end = ((idx + 1) * timePerWord).toFixed(4);
                }});
            }}
            
            function updateHighlight(currentTime) {{
                const followEnabled = localStorage.getItem('agFollowAlongEnabled') !== 'false';
                
                let activeIdx = -1;
                for (let i = wordEls.length - 1; i >= 0; i--) {{
                    const start = parseFloat(wordEls[i].dataset.start || '0');
                    if (currentTime >= start - 0.3) {{
                        activeIdx = i;
                        break;
                    }}
                }}
                
                wordEls.forEach((el, idx) => {{
                    el.classList.remove('active', 'past');
                    if (followEnabled) {{
                        if (idx === activeIdx) {{
                            el.classList.add('active');
                        }} else if (idx < activeIdx) {{
                            el.classList.add('past');
                        }}
                    }}
                }});
            }}
            
            // Use setInterval for consistent timing (RAF gets deprioritized in iframes)
            const parentDoc = window.parent.document;
            let boundAudio = null;
            
            function pollAudio() {{
                const audios = parentDoc.querySelectorAll('audio');
                if (audios.length > 0) {{
                    const audio = audios[audios.length - 1];
                    if (isEstimated && audio.duration && !isNaN(audio.duration)) {{
                        setEstimatedTimings(audio.duration);
                    }}
                    if (boundAudio !== audio) {{
                        boundAudio = audio;
                    }}
                    updateHighlight(audio.currentTime);
                }}
            }}
            
            // Consistent 30ms interval — reliable in iframes unlike RAF
            setInterval(pollAudio, 30);
            
            // Also do initial fast binding
            pollAudio();
        </script>
    """, height=container_height)

def trigger_duplicate_error():
    """Triggers the centered animated error overlay with sound."""
    st.markdown(f"<div class='centered-overlay-error' data-salt='{time.time()}'>⚠️ FILE(S) ALREADY EXISTS!</div>", unsafe_allow_html=True)
    sound_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "audio", "universfield-system-error-notice-132470.mp3")
    if os.path.exists(sound_path):
        with open(sound_path, "rb") as f:
            data = f.read()
            b64 = base64.b64encode(data).decode()
            st.markdown(f'<audio autoplay><source src="data:audio/mp3;base64,{b64}" type="audio/mp3"></audio>', unsafe_allow_html=True)

def cancel_processing():
    """Callback to instantly abort processing and clear the UI."""
    # Kick off background cleanup of any already-uploaded Google Drive files from this batch
    uploaded_ids = st.session_state.get("_uploaded_file_ids", [])
    if uploaded_ids:
        import threading
        def cleanup_drive_files(ids_to_delete):
            import subprocess
            import os
            import sys
            try:
                current_dir = os.path.dirname(os.path.abspath(__file__))
                delete_script = os.path.abspath(os.path.join(current_dir, "..", "_000-Basics", "Data-GoogleDrive", "scripts", "delete_from_drive.py"))
                if not os.path.exists(delete_script):
                    return
                for file_id in ids_to_delete:
                    # Run deletion in background
                    subprocess.run([sys.executable, delete_script, "--id", file_id], capture_output=True)
            except Exception as e:
                pass
                
        threading.Thread(target=cleanup_drive_files, args=(list(uploaded_ids),), daemon=True).start()
        st.session_state["_uploaded_file_ids"] = []

    # Since Streamlit reruns on click, this callback runs before the rest of the script.
    # We clear the active triggers so the script doesn't try to auto-run again on reload
    if get_skill_state("auto_open_result"):
        set_skill_state("auto_open_result", False)

def trigger_processing_overlay():
    """Shows a centered processing banner with a dots animation and a cancel button."""
    placeholder = st.empty()
    
    # Pre-load completion sound as base64 for the JS polling callback
    import base64 as _b64
    _sound_b64 = ""
    try:
        _sound_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "audio", "BEEP-Subtle_call_connecte-Elevenlabs.mp3")
        with open(_sound_path, "rb") as _sf:
            _sound_b64 = _b64.b64encode(_sf.read()).decode()
    except Exception:
        pass
    
    with placeholder.container():
        # Hidden marker so other code can detect processing state
        st.markdown("<div class='processing-marker' style='display:none;'></div>", unsafe_allow_html=True)
        # Add the actual interactive button over the HTML overlay
        st.button("CANCEL", on_click=cancel_processing, key=f"cancel_btn_{time.time()}", help="Immediately stop processing and discard results.")
        
        # Inject the overlay directly into document.body via JS to bypass
        # Streamlit's container hierarchy (which breaks position:fixed).
        # The hidden Streamlit CANCEL button above handles the actual callback.
        import streamlit.components.v1 as components  # type: ignore[import-not-found]
        components.html("""
            <script>
                // cache-buster: """ + str(time.time()) + """
                const doc = window.parent.document;
                
                // Clean up stale elements from previous processing runs
                // so the polling timer doesn't immediately self-destruct
                doc.querySelectorAll('.centered-overlay-complete').forEach(el => el.remove());
                const oldCompleteOverlay = doc.getElementById('ag-complete-overlay');
                if (oldCompleteOverlay) oldCompleteOverlay.remove();
                // Clear any previous polling interval and safety timeout
                const pw = doc.defaultView || window.parent;
                if (doc._agCleanupInterval) {
                    pw.clearInterval(doc._agCleanupInterval);
                    doc._agCleanupInterval = null;
                }
                if (doc._agSafetyTimeout) {
                    pw.clearTimeout(doc._agSafetyTimeout);
                    doc._agSafetyTimeout = null;
                }
                
                // Store completion sound for later use by polling callback
                doc._agCompletionSoundB64 = '""" + _sound_b64 + """';
                
                // Hide the Streamlit placeholder container off-screen
                // (NOT zero dimensions, as that prevents button.click() from working)
                const marker = doc.querySelector('.processing-marker');
                if (marker) {
                    const block = marker.closest('div[data-testid="stVerticalBlock"]');
                    if (block) {
                        block.style.cssText = 'position:fixed !important; left:-9999px !important; top:-9999px !important; opacity:0 !important;';
                    }
                }
                
                // Remove any previous overlay (in case of re-render)
                const oldOverlay = doc.getElementById('ag-processing-overlay');
                if (oldOverlay) oldOverlay.remove();
                
                // Create the overlay directly on document.body
                const overlay = document.createElement('div');
                overlay.id = 'ag-processing-overlay';
                
                // Fade out sidebar when processing starts
                const sidebar = doc.querySelector('[data-testid="stSidebar"]');
                if (sidebar) {
                    sidebar.style.transition = 'opacity 1s ease, transform 1s ease';
                    sidebar.style.opacity = '0';
                    sidebar.style.transform = 'translateX(-100%)';
                    setTimeout(function() {
                        sidebar.style.display = 'none';
                    }, 1000);
                }
                // Scroll to top for consistent positioning on subsequent runs
                var containers = [
                    doc.querySelector('[data-testid="stMain"]'),
                    doc.querySelector('[data-testid="stAppViewContainer"]'),
                    doc.documentElement,
                    doc.body
                ];
                for (var ci = 0; ci < containers.length; ci++) {
                    if (containers[ci]) containers[ci].scrollTop = 0;
                }
                try { window.parent.scrollTo(0, 0); } catch(e) {}
                overlay.innerHTML = `
                    <style>
                        #ag-processing-overlay {
                            position: fixed;
                            top: 0; left: 0; width: 100vw; height: 100vh;
                            z-index: 999990;
                            display: flex; align-items: center; justify-content: center;
                            background-color: rgba(0,0,0,0.97);
                            pointer-events: auto;
                        }
                        #ag-processing-box {
                            position: relative;
                            width: 675px;
                            min-height: 380px;
                            background-color: rgba(10, 10, 15, 0.95);
                            border: 2px solid rgba(136, 224, 228, 0.6);
                            border-radius: 15px;
                            box-shadow: 0 0 15px rgba(172, 240, 241, 0.51),
                                        0 0 40px rgba(172, 240, 241, 0.26),
                                        0 0 80px rgba(172, 240, 241, 0.13),
                                        inset 0 0 20px rgba(172, 240, 241, 0.09);
                            backdrop-filter: blur(10px);
                            padding: 30px;
                            padding-bottom: 80px;
                            text-align: center;
                            display: flex;
                            flex-direction: column;
                            align-items: center;
                            justify-content: center;
                        }
                        #ag-processing-box .ag-title {
                            line-height: 1.1; margin-bottom: 12px;
                            font-size: 2.5rem; font-weight: bold;
                            color: #a8ffdb; font-family: sans-serif;
                        }
                        #ag-processing-box .ag-subtitle {
                            font-size: 1.5rem; opacity: 0.8;
                            font-weight: normal; color: #88d4b4;
                        }
                        #ag-processing-box .ag-dots {
                            display: flex; justify-content: center; gap: 10px; margin: 18px 0;
                        }
                        #ag-processing-box .ag-dot {
                            width: 19px; height: 19px;
                            background-color: #acf0f1; border-radius: 50%;
                            animation: agDotPulse 1.4s infinite ease-in-out both;
                        }
                        #ag-processing-box .ag-dot:nth-child(1) { animation-delay: -0.32s; }
                        #ag-processing-box .ag-dot:nth-child(2) { animation-delay: -0.16s; }
                        @keyframes agDotPulse {
                            0%, 80%, 100% { transform: scale(0); }
                            40% { transform: scale(1.0); }
                        }
                        #ag-processing-box .ag-info {
                            font-size: 1.1rem; color: #fff;
                            font-weight: normal; margin-top: 18px;
                        }
                        #ag-cancel-btn {
                            position: absolute; bottom: 18px; right: 18px;
                            background-color: rgba(200,30,30,0.8);
                            color: rgba(255,255,255,0.9);
                            border: 1px solid rgba(255,100,100,0.4);
                            padding: 6px 18px; font-size: 14px;
                            font-weight: bold; letter-spacing: 0.5px;
                            border-radius: 5px; cursor: pointer;
                            transition: all 0.2s ease;
                        }
                        #ag-cancel-btn:hover {
                            background-color: rgba(230,50,50,0.9);
                            border-color: #ffcccc; color: #fff;
                        }
                    </style>
                    <div id="ag-processing-box">
                        <div class="ag-title">
                            PROCESSING!<br>
                            <span class="ag-subtitle">Please stand by!</span>
                        </div>
                        <div class="ag-dots">
                            <div class="ag-dot"></div>
                            <div class="ag-dot"></div>
                            <div class="ag-dot"></div>
                        </div>
                        <div class="ag-info">Depending on file size: Could Take Up to 5 mins.</div>
                        <button id="ag-cancel-btn">CANCEL</button>
                    </div>
                `;
                doc.body.appendChild(overlay);
                
                // Helper: find and click the hidden Streamlit CANCEL button
                function clickStreamlitCancel() {
                    const m = doc.querySelector('.processing-marker');
                    if (m) {
                        const b = m.closest('div[data-testid="stVerticalBlock"]');
                        if (b) {
                            const btns = Array.from(b.querySelectorAll('button'));
                            const cancel = btns.find(x => x.textContent && x.textContent.includes('CANCEL'));
                            if (cancel) {
                                cancel.click();
                                // Also remove the overlay immediately for responsiveness
                                const o = doc.getElementById('ag-processing-overlay');
                                if (o) o.remove();
                            }
                        }
                    }
                }
                
                // Bridge: JS CANCEL button → hidden Streamlit CANCEL button
                doc.getElementById('ag-cancel-btn').addEventListener('click', clickStreamlitCancel);
                
                // ESCAPE key → CANCEL (remove old listener first to avoid stacking)
                if (doc._agEscListener) {
                    doc.removeEventListener('keydown', doc._agEscListener, true);
                }
                doc._agEscListener = function(ev) {
                    if (ev.key === 'Escape' && doc.getElementById('ag-processing-overlay')) {
                        ev.preventDefault();
                        ev.stopPropagation();
                        clickStreamlitCancel();
                    }
                };
                doc.addEventListener('keydown', doc._agEscListener, true);
                
                // Auto-remove overlay when processing completes.
                // CRITICAL: Use parent window's setInterval (not iframe's) because
                // the iframe is moved offscreen and browsers throttle offscreen iframe timers.
                const parentWin = doc.defaultView || window.parent;
                const cleanupInterval = parentWin.setInterval(function() {
                    const markerStillExists = doc.querySelector('.processing-marker');
                    const completeShown = doc.querySelector('.centered-overlay-complete');
                    if (!markerStillExists || completeShown) {
                        const o = doc.getElementById('ag-processing-overlay');
                        if (o) o.remove();
                        parentWin.clearInterval(cleanupInterval);
                        // Clear safety timeout since we completed normally
                        if (doc._agSafetyTimeout) {
                            parentWin.clearTimeout(doc._agSafetyTimeout);
                            doc._agSafetyTimeout = null;
                        }
                        doc._agCleanupInterval = null;
                        
                        // Show COMPLETE overlay (embedded here so it doesn't depend on a separate iframe)
                        if (!doc.getElementById('ag-complete-overlay')) {
                            const cOverlay = document.createElement('div');
                            cOverlay.id = 'ag-complete-overlay';
                            cOverlay.innerHTML = `
                                <style>
                                    #ag-complete-overlay {
                                        position: fixed;
                                        top: 0; left: 0; width: 100vw; height: 100vh;
                                        z-index: 999990;
                                        display: flex; align-items: center; justify-content: center;
                                        background-color: rgba(0,0,0,0.4);
                                        pointer-events: none;
                                        animation: agCompleteFade 2s ease-in-out forwards;
                                    }
                                    #ag-complete-box {
                                        width: 540px;
                                        background-color: rgba(10, 10, 15, 0.95);
                                        border: 2px solid rgba(136, 224, 228, 0.6);
                                        border-radius: 12px;
                                        box-shadow: 0 0 15px rgba(172, 240, 241, 0.3),
                                                    0 0 40px rgba(172, 240, 241, 0.15),
                                                    0 0 80px rgba(172, 240, 241, 0.08);
                                        padding: 30px;
                                        text-align: center;
                                        font-size: 2rem;
                                        font-weight: bold;
                                        color: #a8ffdb;
                                        font-family: sans-serif;
                                    }
                                    @keyframes agCompleteFade {
                                        0% { opacity: 1; }
                                        70% { opacity: 1; }
                                        100% { opacity: 0; }
                                    }
                                </style>
                                <div id="ag-complete-box">COMPLETE!</div>
                            `;
                            doc.body.appendChild(cOverlay);
                            parentWin.setTimeout(function() {
                                var el = doc.getElementById('ag-complete-overlay');
                                if (el) el.remove();
                                
                                // Curtain reveal — fade out to reveal result overlay
                                var curtain = doc.createElement('div');
                                curtain.id = 'ag-scroll-curtain';
                                curtain.style.cssText = 'position:fixed;top:0;left:0;width:100vw;height:100vh;background:#0a0a0f;opacity:1;z-index:999980;transition:opacity 0.6s ease;pointer-events:none;';
                                doc.body.appendChild(curtain);
                                parentWin.setTimeout(function() {
                                    curtain.style.opacity = '0';
                                    parentWin.setTimeout(function() { curtain.remove(); }, 700);
                                    
                                    // Animate sidebar back in alongside curtain drop
                                    const sb = doc.querySelector('[data-testid="stSidebar"]');
                                    if (sb) {
                                        sb.style.display = '';
                                        // Force reflow so display block applies before transitions
                                        void sb.offsetWidth;
                                        sb.style.transition = 'opacity 0.6s ease, transform 0.6s ease';
                                        sb.style.opacity = '1';
                                        sb.style.transform = 'translateX(0)';
                                    }
                                }, 100);
                            }, 300);
                            
                            // Play completion sound if available
                            if (doc._agCompletionSoundB64) {
                                try {
                                    const audio = new Audio('data:audio/mp3;base64,' + doc._agCompletionSoundB64);
                                    audio.volume = 0.5;
                                    audio.play().catch(function() {});
                                } catch(e) {}
                            }
                        }
                    }
                }, 500);
                // Store interval ID so it can be cleared on subsequent processing runs
                doc._agCleanupInterval = cleanupInterval;
                
                // SAFETY TIMEOUT: Auto-remove overlay after 5 minutes in case
                // polling is orphaned (e.g. st.stop() destroys iframe context)
                doc._agSafetyTimeout = parentWin.setTimeout(function() {
                    const staleOverlay = doc.getElementById('ag-processing-overlay');
                    if (staleOverlay) staleOverlay.remove();
                    if (doc._agCleanupInterval) {
                        parentWin.clearInterval(doc._agCleanupInterval);
                        doc._agCleanupInterval = null;
                    }
                    // Restore sidebar if it was hidden
                    const sb = doc.querySelector('[data-testid="stSidebar"]');
                    if (sb) {
                        sb.style.display = '';
                        sb.style.opacity = '1';
                        sb.style.transform = '';
                    }
                }, 300000);
            </script>
        """, height=0)
        
    return placeholder

def trigger_complete_overlay(placeholder):
    """Replaces processing banner with a complete banner that fades out, with sound."""
    if placeholder:
        # Load the completion sound as base64
        import base64 as _b64
        _sound_b64 = ""
        try:
            _sound_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "audio", "BEEP-Subtle_call_connecte-Elevenlabs.mp3")
            with open(_sound_path, "rb") as _sf:
                _sound_b64 = _b64.b64encode(_sf.read()).decode()
        except Exception:
            pass
        
        with placeholder:
            st.markdown("<div class='centered-overlay-complete' style='display:none;'></div>", unsafe_allow_html=True)
            # Inject JS to remove the processing overlay and show COMPLETE on document.body
            import streamlit.components.v1 as components  # type: ignore[import-not-found]
            components.html(f"""
                <script>
                    const doc = window.parent.document;
                    
                    // Remove the processing overlay
                    const procOverlay = doc.getElementById('ag-processing-overlay');
                    if (procOverlay) procOverlay.remove();
                    
                    // Also unhide the Streamlit container (restore its style)
                    const marker = doc.querySelector('.processing-marker');
                    if (marker) {{
                        const block = marker.closest('div[data-testid="stVerticalBlock"]');
                        if (block) block.style.cssText = '';
                    }}
                    
                    // Remove any old complete overlay
                    const oldComplete = doc.getElementById('ag-complete-overlay');
                    if (oldComplete) oldComplete.remove();
                    
                    // Play completion sound
                    const soundB64 = "{_sound_b64}";
                    if (soundB64) {{
                        try {{
                            const audio = new Audio('data:audio/mp3;base64,' + soundB64);
                            audio.volume = 0.5;
                            audio.play().catch(function() {{}});
                        }} catch(e) {{}}
                    }}
                    
                    // Create COMPLETE overlay on document.body
                    const completeOverlay = document.createElement('div');
                    completeOverlay.id = 'ag-complete-overlay';
                    completeOverlay.innerHTML = `
                        <style>
                            #ag-complete-overlay {{
                                position: fixed;
                                top: 0; left: 0; width: 100vw; height: 100vh;
                                z-index: 999990;
                                display: flex; align-items: center; justify-content: center;
                                background-color: rgba(0,0,0,0.97);
                                pointer-events: none;
                                animation: agCompleteFade 2s ease-in-out forwards;
                            }}
                            #ag-complete-box {{
                                width: 675px;
                                min-height: 380px;
                                background-color: rgba(10, 10, 15, 0.95);
                                border: 2px solid rgba(136, 224, 228, 0.6);
                                border-radius: 15px;
                                box-shadow: 0 0 15px rgba(172, 240, 241, 0.51),
                                            0 0 40px rgba(172, 240, 241, 0.26),
                                            0 0 80px rgba(172, 240, 241, 0.13);
                                padding: 30px;
                                text-align: center;
                                font-size: 2.5rem;
                                font-weight: bold;
                                color: #a8ffdb;
                                font-family: sans-serif;
                                display: flex;
                                align-items: center;
                                justify-content: center;
                            }}
                            @keyframes agCompleteFade {{
                                0% {{ opacity: 1; }}
                                70% {{ opacity: 1; }}
                                100% {{ opacity: 0; }}
                            }}
                        </style>
                        <div id="ag-complete-box">COMPLETE!</div>
                    `;
                    doc.body.appendChild(completeOverlay);
                    
                    // Auto-remove COMPLETE overlay, then curtain reveal
                    setTimeout(function() {{
                        var el = doc.getElementById('ag-complete-overlay');
                        if (el) el.remove();
                        
                        // Curtain reveal — fade out to reveal result overlay
                        var curtain = doc.createElement('div');
                        curtain.id = 'ag-scroll-curtain';
                        curtain.style.cssText = 'position:fixed;top:0;left:0;width:100vw;height:100vh;background:#0a0a0f;opacity:1;z-index:999980;transition:opacity 0.6s ease;pointer-events:none;';
                        doc.body.appendChild(curtain);
                        setTimeout(function() {{
                            curtain.style.opacity = '0';
                            setTimeout(function() {{ curtain.remove(); }}, 700);
                        }}, 100);
                    }}, 300);
                </script>
            """, height=0)

# -----------------------------------------------------------------------------
# Helper: Process Audio/Video Transcription
# -----------------------------------------------------------------------------
def process_uploaded_files(file_paths, selected_skill, run_env, base_args_input="", drive_folder="", google_sheet="", share_with="", proc_overlay=None, main_spinner=None):
    """Executes transcription for a list of files and returns the processed data."""
    results = []
    progress_text = st.empty()
    cwd = selected_skill["dir"]
    
    python_cmd = get_python_cmd()
    
    # Generate batch ID for multi-file grouping in Google Sheets
    import datetime as _dt
    batch_id = _dt.datetime.now().strftime("%Y-%m-%d_%H:%M")
    total_words: int = 0
    
    # --- DEBUG: Log all parameters received ---
    import logging
    logging.basicConfig(level=logging.DEBUG)
    _log = logging.getLogger("process_uploaded_files")
    _log.info(f"[DEBUG] base_args_input = {repr(base_args_input)}")
    _log.info(f"[DEBUG] drive_folder = {repr(drive_folder)}")
    _log.info(f"[DEBUG] google_sheet = {repr(google_sheet)}")
    _log.info(f"[DEBUG] share_with = {repr(share_with)}")
    _log.info(f"[DEBUG] cwd = {repr(cwd)}")
    _log.info(f"[DEBUG] python_cmd = {repr(python_cmd)}")
    _log.info(f"[DEBUG] ELEVENLABS_API_KEY in env = {bool(run_env.get('ELEVENLABS_API_KEY'))}")
    _log.info(f"[DEBUG] GEMINI_API_KEY in env = {bool(run_env.get('GEMINI_API_KEY'))}")
    _log.info(f"[DEBUG] file_paths = {file_paths}")
    # --- END DEBUG ---
    
    for i, fp in enumerate(file_paths):
        progress_text.info(f"⏳ PROCESSING file {i+1} of {len(file_paths)}: `{os.path.basename(fp)}`...")
        
        # Build command with UI-provided arguments
        local_args = str(base_args_input).replace("{FILE_1}", str(fp))
        import shlex
        try:
            local_parsed = shlex.split(local_args)
        except Exception:
            local_parsed = []
            
        if "--input" not in local_parsed and "input" not in str(local_args):
            local_parsed = ["--input", fp] + local_parsed
            
        current_cmd = [python_cmd, selected_skill["script"]] + local_parsed
        if drive_folder:
            current_cmd.extend(["--drive-folder", drive_folder])
        if google_sheet:
            current_cmd.extend(["--google-sheet", google_sheet])
            current_cmd.extend(["--batch-id", batch_id])
            current_cmd.extend(["--batch-seq", str(i + 1)])
        if share_with:
            current_cmd.extend(["--share-with", share_with])
        
        # --- DEBUG: Log exact command ---
        _log.info(f"[DEBUG] FULL CMD: {current_cmd}")
        # --- END DEBUG ---
        
        # Stream the output line by line into the UI
        # Use a new process group so we can kill any sub-subprocesses (like Drive/Sheets uploads)
        process = subprocess.Popen(
            current_cmd, cwd=cwd, env=run_env,
            stdout=subprocess.PIPE, stderr=subprocess.STDOUT,
            text=True, bufsize=1, universal_newlines=True
        )
        
        full_output: list[str] = []
        output_placeholder = st.empty()
        
        try:
            stdout_stream = process.stdout
            if stdout_stream is not None:
                for line in iter(stdout_stream.readline, ""):
                    full_output.append(str(line))
                    # Update UI with the latest 20 lines to keep it snappy
                    output_placeholder.code("".join(full_output[-20:])) # type: ignore
                    
            returncode = process.wait()
        finally:
            if process.poll() is None:
                import signal
                try:
                    process.terminate()
                    process.wait(timeout=2)
                except (subprocess.TimeoutExpired, ProcessLookupError):
                    try:
                        process.kill()
                    except ProcessLookupError:
                        pass
        
        # Combine everything for legacy parsing logic
        final_text = "".join(full_output)
        
        # --- DEBUG: Log subprocess result ---
        _log.info(f"[DEBUG] returncode = {returncode}")
        _dbg_out: str = repr(final_text)
        _log.info(f"[DEBUG] combined output = {_dbg_out[:500]}")  # type: ignore[index]
        # --- END DEBUG ---
        
        transcript = ""
        if returncode == 0:
            ignore_prefixes = (
                "Transcribing:", "Saved:", "Usage:",
                # Google Drive/Sheet machine-readable lines (toggle ON output)
                "FolderID:", "FileID:", "SheetID:", "Link:", "ID:",
                "Auto-uploading", "Logging results", "Converting",
                "Creating folder:", "Uploading ", "Warning:",
                "Sharing file with", "Sharing '",
                "✅ Upload", "✅ Appended", "🔗 Link:",
                # Text2Speech backend log lines
                "Initializing ElevenLabs", "Generating audio",
                "Saving to ", "**Backend", "____", "───",
                # Alignment sidecar log
                "Alignment saved",
            )
            lines = [l for l in final_text.splitlines() if not l.startswith(ignore_prefixes)]
            transcript = "\n".join(lines).strip()
            
            # Track words for batch summary
            total_words += len(transcript.split()) # type: ignore[operator]
            
            # Capture Google IDs from combined output for badge direct-links
            _machine_prefixes = ("Usage:", "Link:", "FolderID:", "SheetID:", "FileID:", "Auto-uploading", "Logging results", "Converting", "Warning:")
            for line in final_text.splitlines():
                if line.startswith("FolderID:"):
                    st.session_state["_google_folder_id"] = line.split(":", 1)[1].strip()
                elif line.startswith("SheetID:"):
                    st.session_state["_google_sheet_id"] = line.split(":", 1)[1].strip()
                elif line.startswith("FileID:"):
                    if "_uploaded_file_ids" not in st.session_state:
                        st.session_state["_uploaded_file_ids"] = []
                    f_id = line.split(":", 1)[1].strip()
                    if f_id not in st.session_state["_uploaded_file_ids"]:
                        st.session_state["_uploaded_file_ids"].append(f_id)
            
            # Show non-machine lines as warnings (but filter out machine-readable ones)
            # Since stdout and stderr are combined, this heuristic might catch script logs.
            # We'll just append the whole cleaned transcript.
            
            # Extract usage — audio_transcribe.py writes Usage: to stderr; text2speech.py uses stdout
            usage_line = next((l for l in final_text.splitlines() if l.startswith("Usage:")), None)
            if usage_line:
                transcript += f"\n\n**Statistics:** {usage_line.split(':', 1)[-1].strip()}"
            
            # Load word-level alignment JSON for synchronized text highlighting (STT)
            import json as _json_stt
            stt_alignment = None
            stt_align_path: str = os.path.splitext(fp)[0] + ".alignment.json"
            if os.path.exists(stt_align_path):
                try:
                    with open(stt_align_path, "r", encoding="utf-8") as alf:
                        stt_alignment = _json_stt.load(alf)
                except Exception:
                    pass
            
            results.append({
                "name": os.path.basename(fp),
                "original_name": os.path.basename(fp),
                "bytes": open(fp, "rb").read() if os.path.exists(fp) else b"",
                "transcript": transcript,
                "alignment": stt_alignment,
            })
        else:
            if "__ANTIGRAVITY_API_QUOTA_EXCEEDED__" in final_text:
                if proc_overlay:
                    proc_overlay.empty()
                if main_spinner:
                    main_spinner.empty()
                progress_text.empty()
                
                # Try to extract detailed message
                import re
                quota_msg = ""
                
                if "kie" in selected_skill["basename"].lower():
                    # Extract the custom Kie.ai msg with bracketed code
                    match = re.search(r"Kie\.ai (?:Upload |Polling )?Error \[(.*?)\]: (.*)(?:\n|$)", final_text)
                    if match:
                        err_code = match.group(1)
                        err_msg = match.group(2).strip()
                        quota_msg = f"\n\n**Details [{err_code}]:** {err_msg}"
                    else:
                        # Fallback for old/unparsed format
                        match_old = re.search(r"Kie\.ai Error: (.*?)(?:\n|$)", final_text)
                        if match_old:
                            quota_msg = f"\n\n**Details:** {match_old.group(1).strip()}"
                    
                    st.error(f"⚠️ **DENIED!** You have reached the maximum usage allowed by your **Kie.ai** balance.{quota_msg}")

                else:
                    # Standard ElevenLabs
                    match = re.search(r"['\"]message['\"]:\s*['\"](.*?)['\"]", final_text)
                    if match:
                        quota_msg = f"\n\n**Usage stats:** {match.group(1)}"
                    st.error(f"⚠️ **DENIED!** You have reached the maximum usage allowed by your **ElevenLabs** active subscription/plan. Please upgrade your plan or wait for the quota to reset.{quota_msg}")
                st.stop()
            
            st.error("Execution Error")
            st.code(final_text)
            continue
    
    progress_text.empty()
    
    # Append batch summary row if files were processed to a Google Sheet
    if google_sheet and len(results) > 0:
        try:
            root_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
            sheet_script = os.path.join(root_dir, "_000-Basics", "Data-GoogleSheet", "scripts", "append_to_sheet.py")
            if os.path.exists(sheet_script):
                summary_cmd = [python_cmd, sheet_script, "--title", google_sheet,
                               "--batch-id", batch_id, "--batch-summary",
                               "--data", f"{len(results)} files", f"{total_words} total words"]
                if share_with:
                    summary_cmd.extend(["--share-with", share_with])
                subprocess.run(summary_cmd, cwd=cwd, capture_output=True, text=True, env=run_env)
        except Exception:
            pass  # Summary row is non-critical
    
    return results

def process_tts_files(file_paths, selected_skill, run_env, base_args_input="", drive_folder="", google_sheet="", share_with="", proc_overlay=None, main_spinner=None):
    """Executes Text2Speech for a list of document files and returns the processed audio data."""
    results = []
    progress_text = st.empty()
    cwd = selected_skill["dir"]
    
    python_cmd = get_python_cmd()
    
    # Generate batch ID for multi-file grouping in Google Sheets
    import datetime as _dt
    batch_id = _dt.datetime.now().strftime("%Y-%m-%d_%H:%M")
    total_words: int = 0
    
    for i, fp in enumerate(file_paths):
        original_name = os.path.basename(fp)
        progress_text.info(f"⏳ PROCESSING document {i+1} of {len(file_paths)}: `{original_name}`...")
        
        local_args = str(base_args_input).replace("{FILE_1}", str(fp))
        import shlex
        try:
            local_parsed = shlex.split(local_args)
        except Exception:
            local_parsed = ["--input", fp]
            
        current_cmd = [python_cmd, selected_skill["script"]] + local_parsed
        if drive_folder:
            current_cmd.extend(["--drive-folder", drive_folder])
        if google_sheet:
            current_cmd.extend(["--google-sheet", google_sheet])
            current_cmd.extend(["--batch-id", batch_id])
            current_cmd.extend(["--batch-seq", str(i + 1)])
        if share_with:
            current_cmd.extend(["--share-with", share_with])
            
        process = subprocess.Popen(current_cmd, cwd=cwd, env=run_env, stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True)
        try:
            stdout_data, stderr_data = process.communicate()
            res = subprocess.CompletedProcess(process.args, process.returncode, stdout_data, stderr_data)
        finally:
            if process.poll() is None:
                process.terminate()
                try:
                    process.wait(timeout=2)
                except subprocess.TimeoutExpired:
                    process.kill()
        
        if res.returncode == 0:
            saved_path = None
            for line in res.stdout.splitlines():
                if "Saved:" in line:
                    saved_path = line.split("Saved:")[1].strip()
                    break
            
            if saved_path:
                saved_path_str: str = str(saved_path)
                full_saved_path = saved_path_str if os.path.isabs(saved_path_str) else os.path.join(cwd, saved_path_str)
                if os.path.exists(str(full_saved_path)):
                    with open(str(full_saved_path), "rb") as af:
                        audio_bytes = af.read()
                    
                    # The text2speech script now outputs the parsed text alongside the audio file
                    txt_path: str = os.path.splitext(full_saved_path)[0] + ".txt"
                    content_preview = ""
                    if os.path.exists(txt_path):
                        try:
                            with open(txt_path, "r", encoding="utf-8") as tf:
                                content_preview = tf.read(5000)
                                if len(content_preview) == 5000:
                                    content_preview += "...\n\n[Preview truncated due to length. Navigate or download for full context]"
                        except:
                            content_preview = f"Narrated Document: {original_name}"
                    else:
                        content_preview = f"Narrated Document: {original_name}"
                    
                    # Load word-level alignment JSON for synchronized text highlighting
                    import json as _json_tts
                    alignment_data = None
                    align_path: str = os.path.splitext(full_saved_path)[0] + ".alignment.json"
                    if os.path.exists(align_path):
                        try:
                            with open(align_path, "r", encoding="utf-8") as alf:
                                alignment_data = _json_tts.load(alf)
                        except Exception:
                            pass
                    
                    # Extract usage (suppress for manual TTS input per user request)
                    usage_line = next((l for l in res.stdout.splitlines() if l.startswith("Usage:")), None)
                    if usage_line:
                        is_manual_input = (os.path.basename(fp) == "input_text.txt")
                        if not is_manual_input:
                            content_preview += f"\n\n**Statistics:** {usage_line.split(':', 1)[-1].strip()}"

                    # Capture Google IDs from stderr for badge direct-links
                    _tts_machine_prefixes = ("Usage:", "Link:", "FolderID:", "SheetID:", "FileID:", "Auto-uploading", "Logging results", "Converting", "Warning:", "Initializing ElevenLabs", "Generating audio", "Saving to ", "Sharing file", "Sharing '", "✅ Upload", "✅ Appended", "🔗 Link:", "ID:", "Alignment saved")
                    for stderr_line in res.stderr.splitlines():
                        if stderr_line.startswith("FolderID:"):
                            st.session_state["_google_folder_id"] = stderr_line.split(":", 1)[1].strip()
                        elif stderr_line.startswith("SheetID:"):
                            st.session_state["_google_sheet_id"] = stderr_line.split(":", 1)[1].strip()

                    # Surface stderr logs as well (filter out machine-readable lines)
                    err_lines = [l for l in res.stderr.strip().splitlines() if not any(l.startswith(p) for p in _tts_machine_prefixes)]
                    if err_lines:
                        content_preview += "\n\n**Backend Logs/Warnings:**\n" + "\n".join(err_lines).strip()

                    results.append({
                        "name": os.path.basename(full_saved_path),
                        "original_name": original_name,
                        "bytes": audio_bytes,
                        "transcript": content_preview,
                        "content_preview": content_preview,
                        "alignment": alignment_data
                    })
                    # Track words for batch summary (use raw content, not formatted preview)
                    total_words += len(content_preview.split()) # type: ignore[operator]
                else:
                    st.error(f"Could not find output file for {original_name}: {saved_path}")
            else:
                 st.error(f"Could not parse 'Saved:' path from output for {original_name}")
        else:
            if "__ANTIGRAVITY_API_QUOTA_EXCEEDED__" in res.stderr:
                if proc_overlay:
                    proc_overlay.empty()
                if main_spinner:
                    main_spinner.empty()
                progress_text.empty()
                
                import re
                quota_msg = ""
                
                if "kie" in selected_skill["basename"].lower():
                    # Extract the custom Kie.ai msg with bracketed code
                    match = re.search(r"Kie\.ai (?:Upload |Polling )?Error \[(.*?)\]: (.*)(?:\n|$)", res.stderr)
                    if match:
                        err_code = match.group(1)
                        err_msg = match.group(2).strip()
                        quota_msg = f"\n\n**Details [{err_code}]:** {err_msg}"
                    else:
                        # Fallback
                        match_old = re.search(r"Kie\.ai Error: (.*?)(?:\n|$)", res.stderr)
                        if match_old:
                            quota_msg = f"\n\n**Details:** {match_old.group(1).strip()}"
                    st.error(f"⚠️ **DENIED!** You have reached the maximum usage allowed by your **Kie.ai** balance.{quota_msg}")

                else:
                    # Standard ElevenLabs
                    match = re.search(r"['\"]message['\"]:\s*['\"](.*?)['\"]", res.stderr)
                    if match:
                        quota_msg = f"\n\n**Usage stats:** {match.group(1)}"
                    st.error(f"⚠️ **DENIED!** You have reached the maximum usage allowed by your **ElevenLabs** active subscription/plan. Please upgrade your plan or wait for the quota to reset.{quota_msg}")
                st.stop()
            st.error("Execution Error")
            st.code(res.stderr)
    
    progress_text.empty()
    
    # Append batch summary row if files were processed to a Google Sheet
    if google_sheet and len(results) > 0:
        try:
            root_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
            sheet_script = os.path.join(root_dir, "_000-Basics", "Data-GoogleSheet", "scripts", "append_to_sheet.py")
            if os.path.exists(sheet_script):
                summary_cmd = [python_cmd, sheet_script, "--title", google_sheet,
                               "--batch-id", batch_id, "--batch-summary",
                               "--data", f"{len(results)} files", f"{total_words} total words"]
                if share_with:
                    summary_cmd.extend(["--share-with", share_with])
                subprocess.run(summary_cmd, cwd=cwd, capture_output=True, text=True, env=run_env)
        except Exception:
            pass  # Summary row is non-critical
    
    return results


st.set_page_config(page_title="Antigravity Skills", page_icon="🚀", layout="wide")
load_css() # Global CSS load to ensure banners work on main page

# -----------------------------------------------------------------------------
# Helper: Generate PDF via Convtr-PlainTxt2PDF Skill
# -----------------------------------------------------------------------------
@st.cache_data(show_spinner=False, max_entries=50)
def generate_pdf_from_text(text: str) -> bytes:
    """Uses the PlainTxt2PDF skill to convert text to PDF bytes."""
    root_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    skill_script = os.path.join(root_dir, "_000-Basics", "Convtr-PlainTxt2PDF", "scripts", "plain_txt2pdf.py")
    
    # Fallback to _100-TBD if not found in _000-Basics
    if not os.path.exists(skill_script):
        skill_script = os.path.join(root_dir, "_100-TBD", "Convtr-PlainTxt2PDF", "scripts", "plain_txt2pdf.py")
    
    if not os.path.exists(skill_script):
        return b""
        
    import tempfile
    with tempfile.TemporaryDirectory() as tmpdir:
        input_txt = os.path.join(tmpdir, "temp_transcript.txt")
        output_pdf = os.path.join(tmpdir, "temp_transcript.pdf")
        
        with open(input_txt, "w", encoding="utf-8") as f:
            f.write(text)
            
        python_cmd = get_python_cmd()
        
        cmd = [python_cmd, skill_script, "--input", input_txt, "--output", output_pdf, "--font-size", "13"]
        res = subprocess.run(cmd, capture_output=True, text=True)
        
        if os.path.exists(output_pdf):
            with open(output_pdf, "rb") as f:
                return f.read()
        else:
            try:
                with open("/tmp/antigravity_pdf_error.log", "a") as errf:
                    errf.write(f"PDF FAILED!\nSTDOUT: {res.stdout}\nSTDERR: {res.stderr}\n")
            except Exception:
                pass
            st.error(f"PDF GENERATION FAILED!\n\nSTDOUT:\n{res.stdout}\n\nSTDERR:\n{res.stderr}")
    return b""

@st.cache_data(show_spinner=False, max_entries=50)
def generate_docx_from_text(text: str) -> bytes:
    """Generates a DOCX file from text returning bytes."""
    import io
    try:
        from docx import Document  # type: ignore[import-not-found]
        from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore[import-not-found]
        from docx.shared import Pt  # type: ignore[import-not-found]
    except ImportError:
        return b""
    
    doc = Document()
    lines = text.split('\n')
    
    for i, line in enumerate(lines):
        p = doc.add_paragraph()
        run = p.add_run(line)
        
        # Check if this is the source file line at the very end
        if line.startswith("Source File:") and i >= len(lines) - 2:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run.font.size = Pt(9) # Keep footer small
        else:
            run.font.size = Pt(13) # Body text ~20% larger than default 11pt
            
    f = io.BytesIO()
    doc.save(f)
    return f.getvalue()

@st.cache_data(show_spinner=False, max_entries=50)
def generate_doc_rtf_from_text(text: str) -> bytes:
    """Generates a basic RTF string for .doc compatibility."""
    # Split to find the source file line
    lines = text.split('\n')
    
    rtf = "{\\rtf1\\ansi\\ansicpg1252\\deff0\\nouicompat\\deflang1033{\\fonttbl{\\f0\\fnil\\fcharset0 Calibri;}}\n"
    rtf += "{\\*\\generator Antigravity;}\\viewkind4\\uc1\n"
    
    for i, line in enumerate(lines):
        escaped_line = line.replace('\\', '\\\\').replace('{', '\\{').replace('}', '\\}')
        
        if line.startswith("Source File:") and i >= len(lines) - 2:
            # \qc = center align, \fs18 = 9pt (~15% smaller than 11pt/22 half-pts)
            rtf += "\\pard\\qc\\sa200\\sl276\\slmult1\\f0\\fs18\\lang9 " + escaped_line + "\\par\n"
        else:
            # \ql = left align, \fs26 = 13pt (20% larger than 11pt)
            rtf += "\\pard\\ql\\sa200\\sl276\\slmult1\\f0\\fs26\\lang9 " + escaped_line + "\\par\n"
            
    rtf += "}"
    return rtf.encode('utf-8')

def read_text_file_preview(path: str) -> str:
    """Reads a text, md, or rtf file for preview purposes."""
    ext = os.path.splitext(path)[1].lower()
    
    # Skip PDF files as they are binary and not human-readable in raw text format
    if ext == ".pdf":
        return ""
    
    # Pre-flight check for RTF files hiding under .txt extension (common with Mac TextEdit)
    is_actually_rtf = ext == ".rtf"
    if not is_actually_rtf and os.path.getsize(path) > 0:
        try:
            with open(path, "r", encoding="utf-8", errors="ignore") as f:
                if f.read(5).startswith("{\\rtf"):
                    is_actually_rtf = True
        except Exception:
            pass

    try:
        # Use native mac util for RTF and legacy DOC files for perfect text matching (via HTML for lists)
        if is_actually_rtf or ext == ".doc":
            import subprocess
            try:
                res = subprocess.run(["textutil", "-convert", "html", "-stdout", path],
                                     capture_output=True, text=True, check=True)
                from html.parser import HTMLParser
                import re
                
                class RTFListParser(HTMLParser):
                    def __init__(self):
                        super().__init__()
                        self.text = []
                        self.in_body = False
                        self.list_stack = []
                        self.ol_counters = []

                    def _add_newline_if_needed(self):
                        if self.text and not self.text[-1].endswith('\n') and not self.text[-1].endswith('\n '):
                            self.text.append('\n')

                    def get_indent(self):
                        level = len(self.list_stack)
                        return "  " * level if level > 0 else ""

                    def handle_starttag(self, tag, attrs):
                        if tag == 'body':
                            self.in_body = True
                        elif tag in ('p', 'div'):
                            self._add_newline_if_needed()
                        elif tag == 'br' and self.in_body:
                            self.text.append('\n')
                        elif tag == 'ul':
                            self.list_stack.append('ul')
                            self.ol_counters.append(0)
                            self._add_newline_if_needed()
                        elif tag == 'ol':
                            self.list_stack.append('ol')
                            self.ol_counters.append(1)
                            self._add_newline_if_needed()
                        elif tag == 'li':
                            self._add_newline_if_needed()
                            indent = self.get_indent()
                            if self.list_stack:
                                list_type = self.list_stack[-1]
                                if list_type == 'ul':
                                    self.text.append(f"{indent}\u2022 ")
                                elif list_type == 'ol':
                                    count = self.ol_counters[-1]
                                    self.text.append(f"{indent}{count}. ")
                                    self.ol_counters[-1] += 1

                    def handle_endtag(self, tag):
                        if tag == 'body':
                            self.in_body = False
                        elif tag in ('p', 'div'):
                            self._add_newline_if_needed()
                        elif tag in ('ul', 'ol'):
                            if self.list_stack:
                                self.list_stack.pop()
                                self.ol_counters.pop()
                            self._add_newline_if_needed()

                    def handle_data(self, data):
                        if self.in_body:
                            if not data.strip():
                                pass
                            else:
                                cleaned = re.sub(r'[\r\n]+', '', data)
                                if cleaned:
                                    self.text.append(cleaned.strip() + " ")
                                
                parser = RTFListParser()
                parser.feed(res.stdout)
                extracted = "".join(parser.text).strip()
                import io as _io
                extracted_str: str = "".join(parser.text).strip()
                return _io.StringIO(re.sub(r'\n{3,}', '\n\n', extracted_str)).read(10000)
            except Exception:
                pass # Fallback below
                
        # Use python-docx for DOCX files
        elif ext == ".docx":
            try:
                import docx  # type: ignore[import-not-found]
                doc = docx.Document(path)
                fullText = []
                for para in doc.paragraphs:
                    fullText.append(para.text)
                import io as _io2
                return _io2.StringIO('\n'.join(fullText)).read(10000)
            except Exception:
                pass # Fallback below

        # Raw read fallback for TXT, MD, Code or failed RTF
        with open(path, "r", encoding="utf-8", errors="replace") as f:
            raw = f.read(10000)
            
        if is_actually_rtf:
            # Simple RTF stripper fallback if textutil failed (Note: strips newlines)
            import re
            text = re.sub(r"\\'[0-9a-fA-F]{2}", "", raw)
            text = re.sub(r"\\[a-z*]+\-?\d*[ ]?", " ", text)
            text = re.sub(r"\\[{}\\]", "", text)
            for _ in range(5): text = re.sub(r"\{[^{}]*\}", " ", text)
            text = re.sub(r"[{}]", "", text)
            text = re.sub(r"\n{3,}", "\n\n", text)
            return text.strip()
            
        return raw

    except Exception as e:
        return f"(Could not generate preview: {e})"
   # Helper: Generate a ZIP file of all transcripts dynamically
@st.cache_data(show_spinner=False, max_entries=50)
def generate_zip_of_all_transcripts(processed_files_list, format_option, include_sources=False):
    import io
    import zipfile
    import os
    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, "a", zipfile.ZIP_DEFLATED, False) as zip_file:
        for i, f in enumerate(processed_files_list):
            fname = os.path.splitext(f["name"])[0]
            
            if format_option.startswith("MP3") or format_option == "MP3 Audio (.mp3)":
                zip_file.writestr(f"{fname}.mp3", f["bytes"])
                continue

            # --- Smart ZIP Pass-through ---
            # If we don't need sources and the file is already the requested format, use its raw bytes directly.
            # This preserves 100% quality and prevents blank pages from empty previews.
            current_ext = os.path.splitext(f["name"])[1].lower()
            is_matching_pdf = format_option.startswith("PDF") and current_ext == ".pdf"
            is_matching_docx = format_option.startswith("DOCX") and current_ext == ".docx"
            
            if not include_sources and (is_matching_pdf or is_matching_docx):
                zip_file.writestr(f["name"], f["bytes"])
                continue

            # Otherwise, it's a document format
            download_text = f.get("content_preview") or ""
            if not download_text.strip():
                transcript = f.get("transcript", "")
                # Filter out success messages if falling back to transcript
                if not any(x in transcript for x in ["Successfully", "Generated:", "✅"]):
                    download_text = transcript
            
            # Prepend Source Header if requested (matching Merge Source style)
            if include_sources:
                header = f"--- Document: {f['name']} ---\n\n"
                download_text = header + download_text
            
            if format_option == "TXT (.txt)" or "TXT" in format_option:
                zip_file.writestr(f"{fname}.txt", download_text)
            elif format_option == "PDF (.pdf)" or "PDF" in format_option:
                pdf_bytes = generate_pdf_from_text(download_text)
                if pdf_bytes:
                    zip_file.writestr(f"{fname}.pdf", pdf_bytes)
                else:
                    zip_file.writestr(f"{fname}_PDF_Generation_Failed.txt", "PDF conversion failed for this file. Please check logs.")
            elif format_option == "DOCX (.docx)" or "DOCX" in format_option:
                docx_bytes = generate_docx_from_text(download_text)
                if docx_bytes:
                    zip_file.writestr(f"{fname}.docx", docx_bytes)
                else:
                    zip_file.writestr(f"{fname}_DOCX_Generation_Failed.txt", "DOCX conversion failed for this file.")
            elif format_option == "DOC (.doc)" or "DOC" in format_option:
                doc_bytes = generate_doc_rtf_from_text(download_text)
                if doc_bytes:
                    zip_file.writestr(f"{fname}.doc", doc_bytes)
                else:
                    zip_file.writestr(f"{fname}_DOC_Generation_Failed.txt", "DOC conversion failed for this file.")
                    
        # Add direct high-fidelity file to ZIP if available and not already included
        direct_file = st.session_state.get("direct_download_file")
        if direct_file:
            zip_file.writestr(direct_file["name"], direct_file["bytes"])
            
    return zip_buffer.getvalue()

# initialization session state tracking 
if "authenticated" not in st.session_state:
    st.session_state["authenticated"] = False

# Helper: get the absolute path to the secret manager script
def _sm_script_path():
    base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    return os.path.join(base_dir, "000C-SKILL_SECURITY-GCSecrtMgr", "scripts", "secret_manager.py")

def _fetch_gcp_secret(secret_name: str) -> str:
    """Fetch a secret value from GCP Secret Manager by calling gcloud directly."""
    try:
        # Setup environment PATH to include the local gcloud SDK if it exists
        env = os.environ.copy()
        base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
        gcloud_bin = os.path.join(base_dir, "google-cloud-sdk", "bin")
        if os.path.isdir(gcloud_bin):
            env["PATH"] = gcloud_bin + ":" + env.get("PATH", "")
        
        # Use gcloud directly — it was confirmed to work in the terminal
        # omitting --project so it uses the active account's project context
        res = subprocess.run(
            ["gcloud", "secrets", "versions", "access", "latest", f"--secret={secret_name}"],
            capture_output=True, text=True, timeout=15, env=env
        )
        if res.returncode != 0:
            # Silently log error for developer debugging
            with open("/tmp/gcp_secret_error.log", "a") as f:
                f.write(f"Error fetching {secret_name}: {res.stderr}\n")
            return ""
        return res.stdout.strip()
    except Exception as e:
        with open("/tmp/gcp_secret_error.log", "a") as f:
            f.write(f"Exception fetching {secret_name}: {str(e)}\n")
        return ""

# -----------------------------------------------------------------------------
# 1. Authentication via GCP Secret Manager
# -----------------------------------------------------------------------------
def check_password():
    """Returns True if the user has entered the correct GCP secret password."""
    load_css() # Apply styles to login screen
    if st.session_state.get("authenticated", False):
        return True

    # --- Center the entire login screen instantly via CSS :has() hook ---
    # We inject a hidden div hook that the CSS relies on. This applies the styles instantly
    # on the very first frame without waiting for Javascript, preventing any layout flash.
    st.markdown('<div id="login-css-hook" style="display:none"></div>', unsafe_allow_html=True)
    st.markdown("""
    <style>
    /* When login-css-hook is present, center everything */
    [data-testid="stAppViewContainer"]:has(#login-css-hook) [data-testid="stMainBlockContainer"] {
        display: flex;
        flex-direction: column;
        align-items: center;
        justify-content: center;
        min-height: 80vh;
        text-align: center;
    }
    /* Center the title */
    [data-testid="stAppViewContainer"]:has(#login-css-hook) [data-testid="stMainBlockContainer"] h1 {
        text-align: center;
        width: 100%;
    }
    /* Center the form submit button */
    [data-testid="stAppViewContainer"]:has(#login-css-hook) [data-testid="stFormSubmitButton"] {
        display: flex;
        justify-content: center;
    }
    }
    [data-testid="stAppViewContainer"]:has(#login-css-hook) [data-testid="stFormSubmitButton"] {
        display: flex;
        justify-content: center;
        width: 100%;
        margin-top: 15px;
    }
    [data-testid="stAppViewContainer"]:has(#login-css-hook) [data-testid="stFormSubmitButton"] button {
        min-width: 160px;
    }
    /* Left-align the "Press Enter to submit" helper text */
    [data-testid="stAppViewContainer"]:has(#login-css-hook) [data-testid="stForm"] [data-testid="stMarkdown"] {
        text-align: left;
    }
    /* Center spinner/status messages on login page instantly to prevent flashes */
    [data-testid="stAppViewContainer"]:has(#login-css-hook) [data-testid="stElementContainer"]:has([data-testid="stSpinner"]),
    [data-testid="stAppViewContainer"]:has(#login-css-hook) .element-container:has([data-testid="stSpinner"]),
    [data-testid="stAppViewContainer"]:has(#login-css-hook) [data-testid="stElementContainer"]:has(.stSpinner),
    [data-testid="stAppViewContainer"]:has(#login-css-hook) .element-container:has(.stSpinner),
    [data-testid="stAppViewContainer"]:has(#login-css-hook) [data-testid="stElementContainer"]:has([data-testid="stStatusWidget"]),
    [data-testid="stAppViewContainer"]:has(#login-css-hook) .element-container:has([data-testid="stStatusWidget"]) {
        display: flex !important;
        justify-content: center !important;
        width: 100% !important;
        margin-top: 2.5rem !important;
    }
    /* Reduce space by 1 line between consecutive spinners/status widgets */
    [data-testid="stAppViewContainer"]:has(#login-css-hook) [data-testid="stElementContainer"]:has([data-testid="stSpinner"]) + [data-testid="stElementContainer"]:has([data-testid="stSpinner"]),
    [data-testid="stAppViewContainer"]:has(#login-css-hook) .element-container:has([data-testid="stSpinner"]) + .element-container:has([data-testid="stSpinner"]),
    [data-testid="stAppViewContainer"]:has(#login-css-hook) [data-testid="stElementContainer"]:has(.stSpinner) + [data-testid="stElementContainer"]:has(.stSpinner),
    [data-testid="stAppViewContainer"]:has(#login-css-hook) .element-container:has(.stSpinner) + .element-container:has(.stSpinner) {
        margin-top: 0.5rem !important;
    }
    [data-testid="stAppViewContainer"]:has(#login-css-hook) [data-testid="stSpinner"],
    [data-testid="stAppViewContainer"]:has(#login-css-hook) [data-testid="stStatusWidget"],
    [data-testid="stAppViewContainer"]:has(#login-css-hook) .stSpinner {
        display: flex;
        justify-content: center;
        text-align: center;
        width: 100%;
    }
    </style>
    """, unsafe_allow_html=True)

    st.markdown("<div class='login-page'>", unsafe_allow_html=True)
    st.markdown("<h2 style='text-align: center; color: #ffffff; margin-bottom: 20px;'>🔒 Antigravity Dashboard</h2>", unsafe_allow_html=True)
    
    with st.form("login_form", clear_on_submit=False, border=False):
        _pw_spacer_l, _pw_col, _pw_spacer_r = st.columns([1, 2, 1])
        with _pw_col:
            password = st.text_input("Password", type="password", placeholder="Enter Password", label_visibility="collapsed")
            st.markdown("<span style='color: #FFE300; font-size: 0.85em;'>Press RETURN to submit</span>", unsafe_allow_html=True)
        
        # Keybind TAB to focus the password input field
        st.components.v1.html(
            """
            <script>
            const doc = window.parent.document;
            if (!doc._tabKeyBound) {
                doc._tabKeyBound = true;
                doc.addEventListener('keydown', function(e) {
                    if (e.key === 'Tab' && e.target.tagName !== 'INPUT' && e.target.tagName !== 'TEXTAREA') {
                        const pwInput = doc.querySelector('input[type="password"]');
                        if (pwInput) {
                            pwInput.focus();
                            e.preventDefault();
                        }
                    }
                });
            }
            </script>
            """,
            height=0
        )
        st.markdown("<br>", unsafe_allow_html=True)
        _btn_l, _btn_c, _btn_r = st.columns([3, 2, 3])
        with _btn_c:
            unlock_clicked = st.form_submit_button("Unlock", use_container_width=True)

    if unlock_clicked:
        _spin_l, _spin_col, _spin_r = st.columns([1, 2, 1])
        with _spin_col:
            with st.spinner("Verifying via Google Cloud Secret Manager..."):
                try:
                    real_secret = _fetch_gcp_secret("DEV-TEST1-LOGIN")
                    if not real_secret:
                        st.error("Failed to fetch secret from GCP. Check your gcloud auth.")
                        return False
                        
                    if password == real_secret:
                        st.session_state["authenticated"] = True
                        
                        # Auto-fetch API keys from GCP secrets at login
                        with st.spinner("Loading API keys from GCP..."):
                            gemini_key = _fetch_gcp_secret("DEV-TEST2-GEMINI")
                            if gemini_key:
                                st.session_state["GEMINI_API_KEY"] = gemini_key
                            elevenlabs_key = _fetch_gcp_secret("DEV-TEST3-11LABS")
                            if elevenlabs_key:
                                st.session_state["ELEVENLABS_API_KEY"] = elevenlabs_key
                            kie_key = _fetch_gcp_secret("DEV-TEST0-KIE")
                            if kie_key:
                                st.session_state["KIE_API_KEY"] = kie_key
                            
                            g_user_email = _fetch_gcp_secret("DEV-TEST5-G_USER")
                            if g_user_email:
                                st.session_state["GCP_USER_EMAIL"] = g_user_email
                        
                        st.rerun()
                    else:
                        st.error("Incorrect password or failed to fetch secret. Ensure you are logged in to gcloud (`gcloud auth login`).")
                except Exception as e:
                    st.error(f"An error occurred: {e}")
    return False

if not check_password():
    st.stop()  # Do not render the rest of the app until authenticated

# --- Stale processing overlay cleanup ---
# If a previous run orphaned the processing overlay (e.g. st.stop() during error),
# remove it immediately on this re-render so the user isn't stuck.
import streamlit.components.v1 as _cleanup_components  # type: ignore[import-not-found]
_cleanup_components.html("""
    <script>
        (function() {
            const doc = window.parent.document;
            const overlay = doc.getElementById('ag-processing-overlay');
            const marker = doc.querySelector('.processing-marker');
            // If overlay exists but no marker, it's stale — remove it
            if (overlay && !marker) {
                overlay.remove();
                // Restore sidebar if it was hidden
                const sb = doc.querySelector('[data-testid="stSidebar"]');
                if (sb) {
                    sb.style.display = '';
                    sb.style.opacity = '1';
                    sb.style.transform = '';
                }
            }
        })();
    </script>
""", height=0)


# -----------------------------------------------------------------------------
# 2. Skill Discovery & Parsing
# -----------------------------------------------------------------------------
def discover_skills():
    """Scans the parent directory recursively for all SKILL.md files and extracts their info."""
    skills = []
    # Build absolute path to the skills root based on this app.py
    root_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    
    # Exclude the dashboard and backup folders from the search
    excludes = ["__000-DASHBOARD-TEST1", "000A_BKUP", ".venv", "__pycache__", ".git"]
    
    for dirpath, dirnames, filenames in os.walk(root_dir):
        # Modification: avoid searching inside excluded directories
        dirnames[:] = [d for d in dirnames if not any(ex in d for ex in excludes)]  # type: ignore[index]
        
        if "SKILL.md" in filenames:
            # Skill ID is now the relative path from root_dir for uniqueness
            skill_rel_path = os.path.relpath(dirpath, root_dir)
            skill_dir_name = os.path.basename(dirpath)
            
            # Category is the top-level directory name relative to root
            rel_parts = skill_rel_path.split(os.sep)
            # If skill is in root, it has no specific category folder
            category = rel_parts[0] if len(rel_parts) > 1 else "General"
            
            skill_md_path = os.path.join(dirpath, "SKILL.md")
            
            try:
                # Parse the YAML frontmatter
                with open(skill_md_path, "r", encoding="utf-8") as f:
                    content = f.read()
                    
                parts = content.split("---")
                if len(parts) >= 3:
                    frontmatter = yaml.safe_load(parts[1])
                    # The "name" from YAML is now a "display_title" (subtitle)
                    # The primary 'name' used in the dashboard is the folder basename
                    display_title = frontmatter.get("name", skill_dir_name)
                    desc = frontmatter.get("description", "No description provided.")
                else:
                    display_title = skill_dir_name
                    desc = "No frontmatter found in SKILL.md"
                    
                # Unified Antigravity naming: Primary name = folder basename
                name = skill_dir_name
                    
                # Look for the main python script
                # 1. Check if frontmatter specifies a script
                fm_script = frontmatter.get("script")
                if fm_script:
                    main_script = os.path.join(dirpath, fm_script)
                else:
                    # 2. Fallback to glob (take the first one)
                    scripts = glob.glob(os.path.join(dirpath, "scripts", "*.py"))
                    main_script = scripts[0] if scripts else None
                
                skills.append({
                    "id": skill_rel_path,
                    "basename": skill_dir_name,
                    "name": name,
                    "display_title": display_title,
                    "category": category,
                    "desc": desc,
                    "dir": dirpath,
                    "script": main_script
                })
                
                # Once a skill is found, we stop walking deeper into THIS specific directory tree
                dirnames[:] = []  # type: ignore[index]
                
            except Exception:
                # Silently fail discovery for malformed/missing SKILL.md
                pass
                
    return skills


# -----------------------------------------------------------------------------
# 3. Main Dashboard UI
# -----------------------------------------------------------------------------
skills = discover_skills()

st.sidebar.markdown(
    '<div style="color: white; font-size: 1.75rem; font-weight: 700; '
    'line-height: 1.2; padding: 0.75rem 0 0.25rem 0;">🚀 Antigravity Skills</div>',
    unsafe_allow_html=True
)

# --- Theme Toggle (Light/Dark) on same line as skills count ---
if "theme_mode" not in st.session_state:
    st.session_state["theme_mode"] = "dark"

# Read current toggle state (from previous run) to sync icon
_current_light = st.session_state.get("_theme_toggle", False)
st.session_state["theme_mode"] = "light" if _current_light else "dark"

_info_col, _icon_col, _toggle_col, _tip_col = st.sidebar.columns([4, 1.2, 1, 0.5])
with _info_col:
    st.markdown(f"<div style='text-align: center; color: white;'><b>{len(skills)} skills loaded</b></div>", unsafe_allow_html=True)
with _icon_col:
    if _current_light:
        st.components.v1.html(
            '<div style="font-size:30px; text-align:center; margin-left:-5px; line-height:1; margin-top:2px;">☀️</div>',
            height=35
        )
    else:
        st.components.v1.html(
            '<div style="font-size:30px; text-align:center; color:#5b9bd5; margin-left:-5px; '
            'transform:scaleX(-1); line-height:1; margin-top:2px;">☾</div>',
            height=35
        )
with _toggle_col:
    st.toggle("", value=_current_light, key="_theme_toggle", label_visibility="collapsed")
with _tip_col:
    st.markdown("""<div style='margin-top:5px;'>
<style>
.info-tip { position:relative; display:inline-block; font-size:16px; color:white; cursor:pointer; }
.info-tip:hover::after {
    content:'Toggle Light / Dark mode';
    position:absolute; top:125%; right:0;
    background:#333; color:#fff; padding:6px 12px; border-radius:6px;
    font-size:14px; white-space:nowrap; z-index:999;
}
</style>
<span class="info-tip">ⓘ</span></div>""", unsafe_allow_html=True)

# Style the theme toggle sun icon larger (25px)
st.markdown("""
<style>
[data-testid="stSidebar"] div[data-testid="stToggle"][class*=""] label p {
    font-size: 25px !important;
    line-height: 1 !important;
}
</style>
""", unsafe_allow_html=True)

# Spacer between toggle row and search
st.sidebar.markdown("<div style='margin-bottom: 11px;'></div>", unsafe_allow_html=True)

if st.session_state["theme_mode"] == "light":
    st.markdown("""
    <style>
    /* ── Light Mode Overrides ── */
    /* Main background */
    .stApp, [data-testid="stAppViewContainer"], section[data-testid="stMain"] {
        background-color: #ffffff !important;
        color: #1a1a1a !important;
    }
    /* Top header bar */
    header[data-testid="stHeader"] {
        background-color: #ffffff !important;
    }

    /* ── SIDEBAR ── */
    /* Sidebar background */
    [data-testid="stSidebar"], [data-testid="stSidebar"] > div {
        background-color: #666668 !important;
    }
    /* Sidebar text — white for readability on dark grey bg */
    [data-testid="stSidebar"] h1,
    [data-testid="stSidebar"] h2,
    [data-testid="stSidebar"] h3,
    [data-testid="stSidebar"] > div > div > div > div > p,
    [data-testid="stSidebar"] label,
    [data-testid="stSidebar"] [data-testid="stMarkdownContainer"] p,
    [data-testid="stSidebar"] [data-testid="stMarkdownContainer"] span,
    [data-testid="stSidebar"] [data-testid="stMarkdownContainer"] strong {
        color: #ffffff !important;
    }
    /* Sidebar buttons — all consistent lighter grey */
    [data-testid="stSidebar"] .stButton > button {
        background-color: #828589 !important;
        color: #ffffff !important;
        border-color: #828589 !important;
    }
    [data-testid="stSidebar"] .stButton > button p,
    [data-testid="stSidebar"] .stButton > button span {
        color: #ffffff !important;
    }
    [data-testid="stSidebar"] .stButton > button:hover {
        background-color: #595959 !important;
        border-color: #595959 !important;
    }
    /* Sidebar expander — header dark grey, content area lighter */
    [data-testid="stSidebar"] [data-testid="stExpander"] {
        background-color: #5d5e61 !important;
        border-color: #606367 !important;
        border-radius: 8px !important;
        overflow: hidden;
    }
    [data-testid="stSidebar"] [data-testid="stExpander"] summary {
        background-color: #43484e !important;
        color: #ffffff !important;
    }
    [data-testid="stSidebar"] [data-testid="stExpander"] summary span,
    [data-testid="stSidebar"] [data-testid="stExpander"] summary p,
    [data-testid="stSidebar"] [data-testid="stExpander"] summary svg {
        color: #ffffff !important;
        fill: #ffffff !important;
    }
    /* Sidebar search input */
    [data-testid="stSidebar"] input {
        background-color: #ffffff !important;
        color: #1a1a1a !important;
        border-color: #999 !important;
    }
    /* Sidebar horizontal rules */
    [data-testid="stSidebar"] hr {
        border-color: #999 !important;
    }
    /* Sidebar close/collapse button */
    [data-testid="stSidebar"] button[kind="header"] {
        color: #1a1a1a !important;
    }
    /* Sidebar success/warning alerts — darker for readability */
    [data-testid="stSidebar"] [data-testid="stAlert"] {
        background-color: rgba(15, 130, 60, 0.45) !important;
        border-color: #0a7a3a !important;
    }
    [data-testid="stSidebar"] [data-testid="stAlert"] p,
    [data-testid="stSidebar"] [data-testid="stAlert"] span,
    [data-testid="stSidebar"] [data-testid="stAlert"] [data-testid="stMarkdownContainer"] p,
    [data-testid="stSidebar"] [data-testid="stAlert"] [data-testid="stMarkdownContainer"] span,
    [data-testid="stSidebar"] [data-testid="stAlert"] * {
        color: #ffffff !important;
        font-weight: 600 !important;
    }

    /* ── MAIN CONTENT ── */
    /* Text */
    .stApp h1, .stApp h2, .stApp h3, .stApp h4, .stApp p, .stApp label, .stApp li {
        color: #1a1a1a !important;
    }
    /* Input fields & text areas */
    .stApp input, .stApp textarea {
        background-color: #f7f7f8 !important;
        color: #1a1a1a !important;
        border-color: #d0d0d0 !important;
    }
    /* Main content expanders */
    section[data-testid="stMain"] [data-testid="stExpander"] {
        background-color: #f0f1f3 !important;
        border-color: #d0d0d0 !important;
    }
    section[data-testid="stMain"] [data-testid="stExpander"] summary span {
        color: #1a1a1a !important;
    }
    /* File uploader */
    .stApp [data-testid="stFileUploader"] section {
        background-color: #f7f7f8 !important;
        border-color: #d0d0d0 !important;
    }
    .stApp [data-testid="stFileUploader"] section small,
    .stApp [data-testid="stFileUploader"] section span {
        color: #555 !important;
    }
    /* Select boxes / dropdowns */
    .stApp [data-baseweb="select"],
    .stApp [data-baseweb="select"] div {
        background-color: #f7f7f8 !important;
        color: #1a1a1a !important;
    }
    /* Markdown containers */
    .stApp [data-testid="stMarkdownContainer"] {
        color: #1a1a1a !important;
    }
    /* Transcript box */
    .transcript-box {
        background-color: #f0f0f0 !important;
        color: #1a1a1a !important;
        border-color: #cccccc !important;
    }
    /* Info/alert boxes — ensure text readable */
    .stApp [data-testid="stAlert"] p {
        color: inherit !important;
    }
    /* Toggle visibility in light mode — use filter to darken at render layer */
    section[data-testid="stMain"] [data-testid="stToggle"] [role="checkbox"][aria-checked="false"] {
        filter: invert(0.5) contrast(1.3) !important;
    }
    </style>
    """, unsafe_allow_html=True)

# Add Search Box
query_params = st.query_params
default_search = query_params.get("search", "")
search_query = st.sidebar.text_input("🔍 Search skills...", value=default_search).lower()

# Sync search to query params
if search_query:
    st.query_params["search"] = search_query
elif "search" in query_params:
    # Clear search if input is empty
    st.query_params.pop("search")

# Filter skills dynamically
if search_query:
    display_skills = [
        s for s in skills 
        if (s.get('name') and search_query in str(s.get('name')).lower()) or 
           (s.get('desc') and search_query in str(s.get('desc')).lower()) or
           (s.get('category') and search_query in str(s.get('category')).lower())
    ]
else:
    display_skills = skills

# Show API key status
if st.session_state.get("GEMINI_API_KEY"):
    st.sidebar.success("🔑 GEMINI_API_KEY loaded from GCP")
else:
    st.sidebar.warning("⚠️ GEMINI_API_KEY not found (DEV-TEST2-GEMINI)")
if st.session_state.get("ELEVENLABS_API_KEY"):
    st.sidebar.success("🔑 ELEVENLABS_API_KEY loaded from GCP")
else:
    st.sidebar.warning("⚠️ ELEVENLABS_API_KEY not found (DEV-TEST3-11LABS)")

if not skills:
    st.error("No skills found. Please ensure SKILL.md files exist in the parent directories.")
    st.stop()
    
# Sort final list by display name alphabetically so it makes sense to the user
display_skills = sorted(display_skills, key=lambda s: (s.get('name') or '').lower())

if not display_skills:
    st.sidebar.warning(f"No skills match '{search_query}'.")
    st.warning("Please clear the search to see all available skills.")
    st.stop()

import json

# Setup Recent Skills
RECENT_SKILLS_FILE = os.path.join(os.path.dirname(__file__), "recent_skills.json")

def load_recent_skills():
    if os.path.exists(RECENT_SKILLS_FILE):
        try:
            with open(RECENT_SKILLS_FILE, "r") as f:
                return json.load(f)
        except Exception:
            return []
    return []

def save_recent_skill(skill_id):
    recents = load_recent_skills()
    if skill_id in recents:
        recents.remove(skill_id)
    recents.insert(0, skill_id)
    import io as _io_r
    _joined = " ".join(str(r) for r in recents)
    recents = _joined.split()[:4]  # type: ignore[misc]  # Keep top 4
    try:
        with open(RECENT_SKILLS_FILE, "w") as f:
            json.dump(recents, f)
    except Exception:
        pass
    return recents

# Helper: Skill-Specific Session State Namespacing
def get_skill_state(key, default=None):
    """Returns a namespaced session state key for the current skill."""
    if "selected_skill_id" not in st.session_state:
        return default
    ns_key = f"{st.session_state.selected_skill_id}_{key}"
    return st.session_state.get(ns_key, default)

def set_skill_state(key, value):
    """Sets a namespaced session state key for the current skill."""
    if "selected_skill_id" not in st.session_state:
        return
    ns_key = f"{st.session_state.selected_skill_id}_{key}"
    st.session_state[ns_key] = value

# Render Recent Skills
recent_ids = load_recent_skills()
valid_recent_skills = [s for id_ in recent_ids for s in skills if s["id"] == id_]

if valid_recent_skills:
    st.sidebar.markdown("### Recent Used:")
    for rs in valid_recent_skills:
        if st.sidebar.button(rs["name"], use_container_width=True, key=f"recent_{rs['id']}"):
            st.session_state.selected_skill_id = rs["id"]
            st.query_params["skill"] = rs["id"] # NEW: Sync on update
            st.rerun()
    st.sidebar.markdown("---")

# Skill selection
# NEW: Check query params for skill persistence
query_params = st.query_params
if "skill" in query_params:
    p_id = query_params["skill"]
    # Check against FULL skills list, not filtered display_skills
    if any(s["id"] == p_id for s in skills):
        st.session_state.selected_skill_id = p_id
    else:
        # Fallback if invalid ID
        default_id = None
        recents = load_recent_skills()
        if recents and any(s["id"] == recents[0] for s in display_skills):
            default_id = recents[0]
        if not default_id:
            default_id = next((s["id"] for s in display_skills if s["basename"] == "AI-LLM-Text2Speech"), 
                              display_skills[0]["id"] if display_skills else None)
        st.session_state.selected_skill_id = default_id
else:
    # Default selection logic: prefer last used
    default_id = None
    recents = load_recent_skills()
    if recents:
        # Verify the requested recent skill actually still exists in loaded skills
        if any(s["id"] == recents[0] for s in display_skills):
            default_id = recents[0]
            
    # Fallback if no valid recent skill was found
    if not default_id:
        default_id = next((s["id"] for s in display_skills if s["basename"] == "AI-LLM-Text2Speech"), 
                          display_skills[0]["id"] if display_skills else None)
    st.session_state.selected_skill_id = default_id

# Sync query params
if st.session_state.get("selected_skill_id"):
    st.query_params["skill"] = st.session_state.selected_skill_id

selected_skill_id = st.session_state.get("selected_skill_id")
selected_skill = next((s for s in skills if s["id"] == selected_skill_id), None)

if selected_skill_id:
    save_recent_skill(selected_skill_id)

# Group skills by category for display
categories = sorted(list(set(s["category"] for s in display_skills)))

for cat in categories:
    cat_skills = [s for s in display_skills if s["category"] == cat]
    if not cat_skills:
        continue
        
    with st.sidebar.expander(f"📁 {cat}", expanded=bool(selected_skill and selected_skill.get("category") == cat)):  # type: ignore[union-attr]
        for s in cat_skills:
            # Highlight selected skill
            is_selected = st.session_state.selected_skill_id == s["id"]
            btn_label = f"🚀 {s['name']}" if is_selected else s["name"]
            if st.button(btn_label, key=f"select_{s['id']}", use_container_width=True):
                st.session_state.selected_skill_id = s["id"]
                st.query_params["skill"] = s["id"] # NEW: Sync on update
                st.rerun()

if not selected_skill:
    st.error("Failed to load the selected skill.")
    st.stop()


if selected_skill is None:
    st.stop()
assert selected_skill is not None

st.title(selected_skill["name"], anchor=False)
if selected_skill.get("display_title") and selected_skill["display_title"] != selected_skill["name"]:
    st.markdown(f"### *{selected_skill['display_title']}*")
st.info(selected_skill["desc"])

if not selected_skill["script"]:
    st.warning("No executable python script found in the 'scripts' folder for this skill.")
    st.stop()
    
st.markdown("---")

is_audio_skill = selected_skill["basename"] in ["AI-LLM-Speech2Text", "AI-LLM-KIE-ElevenLabs-Speech2Text"]
is_tts_skill = selected_skill["basename"] in ["AI-LLM-Text2Speech", "AI-LLM-KIE-ElevenLabs-Text2Speech"]
is_image_skill = selected_skill["basename"] == "AI-LLM-ImageGenerate"
is_embed_skill = selected_skill["basename"] == "AI-LLM-EmbedText"
is_rag_skill = selected_skill["basename"] == "AI-LLM-RAGQuery"
is_translate_skill = selected_skill["basename"] == "AI-LLM-TranslateText"
is_pdf_skill = selected_skill["basename"] == "Convtr-PlainTxt2PDF"

# --- Watch Folder Auto-Process (Convtr-PlainTxt2PDF only) ---
if is_pdf_skill:
    st.subheader("Create Watch Folder", anchor=False)
    
    # --- Persistent config file for watch folder settings ---
    _wf_config_dir = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "_000-Basics", "Convtr-PlainTxt2PDF", "scripts")
    _wf_config_path = os.path.join(_wf_config_dir, ".watch_folder_config.json")
    
    def _wf_load_config() -> dict:
        """Load saved watch folder config from disk."""
        try:
            if os.path.exists(_wf_config_path):
                import json as _json
                with open(_wf_config_path, "r") as f:
                    return _json.load(f)
        except Exception:
            pass
        return {}
    
    def _wf_save_config(path: str, interval_idx: int):
        """Save watch folder config to disk."""
        try:
            import json as _json
            os.makedirs(os.path.dirname(_wf_config_path), exist_ok=True)
            with open(_wf_config_path, "w") as f:
                _json.dump({"path": path, "interval_idx": interval_idx}, f)
        except Exception:
            pass
    
    def _wf_clear_config():
        """Clear saved watch folder config and reset session state."""
        try:
            if os.path.exists(_wf_config_path):
                os.remove(_wf_config_path)
        except Exception:
            pass
        st.session_state.pop("_wf_picked_path", None)
        st.session_state.pop("_wf_typed_path", None)
        st.session_state.pop("_watch_interval_idx", None)
        st.session_state.pop("_watch_folder_on", None)
        st.session_state.pop("_watch_folder_last_run", None)
        st.session_state.pop("_wf_picker_proc", None)
        st.session_state.pop("_wf_config_loaded", None)
    
    # Load config on first run (hydrate session state from disk)
    if "_wf_config_loaded" not in st.session_state:
        _saved_cfg = _wf_load_config()
        if _saved_cfg.get("path"):
            st.session_state["_wf_picked_path"] = _saved_cfg["path"]
            st.session_state["_watch_interval_idx"] = _saved_cfg.get("interval_idx", 2)
            st.session_state["_watch_folder_on"] = True
        st.session_state["_wf_config_loaded"] = True
    
    _wf_col_toggle, _wf_col_clear = st.columns([1, 2])
    with _wf_col_toggle:
        watch_folder_on = st.toggle("📂 Watch Folder Auto-Process", value=st.session_state.get("_watch_folder_on", False), key="_watch_folder_on")
    with _wf_col_clear:
        pass  # Clear button moved below
    
    if watch_folder_on:
        # --- Folder picker via macOS native dialog (synchronous in callback) ---
        def _pick_folder_dialog():
            """Launch macOS native folder picker. Blocks callback until user picks or cancels."""
            try:
                result = subprocess.run(
                    ["osascript", "-e",
                     'return POSIX path of (choose folder with prompt "Select or Create Watch Folder")'],
                    capture_output=True, text=True, timeout=300
                )
                if result.returncode == 0:
                    picked = result.stdout.strip()
                    if picked:
                        st.session_state["_wf_picked_path"] = picked
                        # Save to disk immediately so it persists across restarts
                        _wf_save_config(picked, st.session_state.get("_watch_interval_idx", 2))
            except (subprocess.TimeoutExpired, Exception):
                pass
        
        # Merge picked path into the display value
        _wf_display_path = st.session_state.get("_wf_picked_path", "") or st.session_state.get("_wf_typed_path", "")
        
        def _open_folder_in_finder():
            """Open the watch folder in Finder."""
            _path = st.session_state.get("_wf_picked_path", "") or st.session_state.get("_wf_typed_path", "")
            if _path:
                _expanded = os.path.expanduser(_path)
                if os.path.isdir(_expanded):
                    subprocess.Popen(["open", _expanded])
        
        # Row 1: Path input + Polling Interval
        _wf_col_path, _wf_col_interval = st.columns([4, 1.2])
        with _wf_col_path:
            st.markdown('<span style="color: #FF8C00; font-weight: 600;">Watch Folder Path</span>', unsafe_allow_html=True)
            watch_folder_path = st.text_input(
                "Watch Folder Path",
                value=_wf_display_path,
                placeholder="Click Create or paste a path...",
                key="_wf_typed_path",
                label_visibility="collapsed"
            )
            # Inject native title tooltip on the text input for full path on hover
            _tt_path = (_wf_display_path or "").replace("'", "\\'")
            st.components.v1.html(f"""
                <script>
                (function() {{
                    var doc = window.parent.document;
                    var el = doc.querySelector('input[aria-label="Watch Folder Path"]');
                    if (el) {{ el.title = '{_tt_path}'; }}
                }})();
                </script>
            """, height=0)
            # If user typed a path manually, clear the picked path so typed takes precedence
            if watch_folder_path and watch_folder_path != st.session_state.get("_wf_picked_path", ""):
                st.session_state.pop("_wf_picked_path", None)
        with _wf_col_interval:
            # Show spinner only when there are pending files being processed
            _is_actively_polling = bool(st.session_state.get("_wf_pending_count", 0))
            _spinner_html = ' <span class="wf-spinner"></span>' if _is_actively_polling else ""
            st.markdown(f'<span style="font-size: 0.875rem; font-weight: 400;">Polling Interval{_spinner_html}</span>', unsafe_allow_html=True)
            watch_interval_label = st.selectbox(
                "Polling Interval",
                options=["15 minutes", "1 minute", "5 seconds"],
                index=st.session_state.get("_watch_interval_idx", 2),
                key="_watch_interval_select",
                label_visibility="collapsed"
            )
            # Store selected index for persistence
            _interval_options = ["15 minutes", "1 minute", "5 seconds"]
            st.session_state["_watch_interval_idx"] = _interval_options.index(watch_interval_label) if watch_interval_label in _interval_options else 2
        
        # Row 2: Action buttons — CREATE (primary) | OPEN | CLEAR — all in one line
        _wf_effective_path = watch_folder_path or _wf_display_path
        _has_valid_path = bool(_wf_effective_path) and os.path.isdir(os.path.expanduser(_wf_effective_path))
        _wf_folder_name = os.path.basename(os.path.expanduser(_wf_effective_path).rstrip("/")) if _wf_effective_path else "Open"
        _has_config = os.path.exists(_wf_config_path) or st.session_state.get("_wf_picked_path") or st.session_state.get("_wf_typed_path")
        
        _btn_cols = st.columns([1, 1, 1, 2]) if _has_config else st.columns([1, 1, 3])
        with _btn_cols[0]:
            st.button("📁 Create", on_click=_pick_folder_dialog, use_container_width=True, key="_wf_browse_btn", type="primary")
        with _btn_cols[1]:
            st.button(f"📂 {_wf_folder_name}", on_click=_open_folder_in_finder, key="_wf_open_btn", disabled=not _has_valid_path, use_container_width=True)
        if _has_config:
            with _btn_cols[2]:
                st.button("🗑️ Clear", on_click=_wf_clear_config, key="_wf_clear_btn", use_container_width=True)
        
        # Style buttons + spinner animation
        st.markdown("""
            <style>
            /* Spinning animation for active polling */
            @keyframes wfSpin {
                0% { transform: rotate(0deg); }
                100% { transform: rotate(360deg); }
            }
            .wf-spinner {
                display: inline-block;
                width: 10px;
                height: 10px;
                border: 2px solid rgba(0, 200, 83, 0.3);
                border-top: 2px solid #00c853;
                border-radius: 50%;
                margin-left: 6px;
                vertical-align: middle;
                animation: wfSpin 0.8s linear infinite;
            }
            /* Watch folder button font size */
            div[data-testid="stVerticalBlock"] button[kind="secondary"] {
                font-size: 0.85em !important;
            }
            </style>
        """, unsafe_allow_html=True)
        
        # Map label to milliseconds for JS timer
        _interval_ms_map = {"15 minutes": 900000, "1 minute": 60000, "5 seconds": 5000}
        _watch_interval_ms = _interval_ms_map.get(watch_interval_label, 60000)
        
        # Show estimated next poll time (skip for 5s interval)
        if _is_actively_polling and watch_interval_label != "5 seconds":
            import datetime as _wf_dt
            _last_run_ts = st.session_state.get("_watch_folder_last_run", 0)
            _interval_secs = _watch_interval_ms / 1000
            if _last_run_ts:
                _next_poll_ts = _last_run_ts + _interval_secs
                _next_poll_time = _wf_dt.datetime.fromtimestamp(_next_poll_ts).strftime("%-I:%M %p")
                st.caption(f"⏱️ Next poll ~{_next_poll_time}")
            else:
                st.caption("⏱️ Next poll on first cycle")
        
        # Expand and validate the folder path
        _effective_path = watch_folder_path or _wf_display_path
        _resolved_watch_path = os.path.expanduser(_effective_path) if _effective_path else ""
        
        if _resolved_watch_path:
            if not os.path.isdir(_resolved_watch_path):
                try:
                    os.makedirs(_resolved_watch_path, exist_ok=True)
                    st.success(f"✅ Created watch folder: `{_resolved_watch_path}`")
                except OSError as e:
                    st.error(f"❌ Could not create folder: {e}")
                    _resolved_watch_path = ""
            
            if _resolved_watch_path:
                # Persist config to disk
                _wf_save_config(_resolved_watch_path, st.session_state.get("_watch_interval_idx", 2))
                
                # Count pending files
                _supported_exts = {".txt", ".rtf", ".doc", ".docx"}
                try:
                    _pending_files = [f for f in os.scandir(_resolved_watch_path) 
                                      if f.is_file() and os.path.splitext(f.name)[1].lower() in _supported_exts]
                except OSError:
                    _pending_files = []
                st.session_state["_wf_pending_count"] = len(_pending_files)
                
                _processed_dir = os.path.join(_resolved_watch_path, "zProcessed")
                _total_processed = 0
                if os.path.isdir(_processed_dir):
                    for _date_dir in os.scandir(_processed_dir):
                        if _date_dir.is_dir():
                            _total_processed += len([f for f in os.scandir(_date_dir.path) if f.is_file()])
                
                # --- Auto-process pending files ---
                import datetime as _wf_dt
                if _pending_files:
                    _last_run_key = "_watch_folder_last_run"
                    _now = _wf_dt.datetime.now().timestamp()
                    _last_run = st.session_state.get(_last_run_key, 0)
                    _interval_secs = _watch_interval_ms / 1000
                    
                    if (_now - _last_run) >= _interval_secs:
                        st.session_state[_last_run_key] = _now
                        
                        # Run the processor
                        _root_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
                        _processor_script = os.path.join(_root_dir, "_000-Basics", "Convtr-PlainTxt2PDF", "scripts", "watch_folder_processor.py")
                        
                        if os.path.exists(_processor_script):
                            _python_cmd = get_python_cmd()
                            _proc_cmd = [_python_cmd, _processor_script, "--folder", _resolved_watch_path]
                            try:
                                _proc_result = subprocess.run(_proc_cmd, capture_output=True, text=True, timeout=300)
                                _full_stdout = _proc_result.stdout.strip()
                                if _full_stdout:
                                    st.session_state["_wf_last_output"] = _full_stdout
                                    st.session_state["_wf_last_output_time"] = _wf_dt.datetime.now().strftime("%-I:%M:%S %p")
                                    
                                    # Feed processed PDFs into the main PROCESSED RESULT display
                                    _today_str = _wf_dt.date.today().strftime("%Y-%m-%d")
                                    _today_dir = os.path.join(_resolved_watch_path, "zProcessed", _today_str)
                                    if os.path.isdir(_today_dir):
                                        _wf_processed_files = []
                                        for _pdf_entry in sorted(os.scandir(_today_dir), key=lambda e: e.name):
                                            if _pdf_entry.is_file() and _pdf_entry.name.lower().endswith(".pdf"):
                                                try:
                                                    with open(_pdf_entry.path, "rb") as _pf:
                                                        _pdf_bytes = _pf.read()
                                                    _wf_processed_files.append({
                                                        "name": _pdf_entry.name,
                                                        "bytes": _pdf_bytes,
                                                        "transcript": f"✅ Converted: {_pdf_entry.name}",
                                                        "content_preview": "",
                                                    })
                                                except (OSError, IOError):
                                                    pass
                                        if _wf_processed_files:
                                            set_skill_state("last_processed_files", _wf_processed_files)
                                            set_skill_state("last_output", _full_stdout)
                                            set_skill_state("file_index", 0)
                                
                                if _proc_result.stderr.strip():
                                    st.session_state["_wf_last_errors"] = _proc_result.stderr.strip()
                                else:
                                    st.session_state.pop("_wf_last_errors", None)
                            except subprocess.TimeoutExpired:
                                st.session_state["_wf_last_errors"] = "Watch folder processing timed out."
                            except Exception as _proc_e:
                                st.session_state["_wf_last_errors"] = f"Watch folder error: {_proc_e}"
                
                with st.expander("DETAILS", expanded=False):
                    st.caption(f"📁 `{_resolved_watch_path}` — **{len(_pending_files)}** pending · **{_total_processed}** processed")
                    
                    # Show last processing output (persistent across reruns)
                    _last_output = st.session_state.get("_wf_last_output", "")
                    if _last_output:
                        _out_time = st.session_state.get("_wf_last_output_time", "")
                        st.success(f"**Converted at {_out_time}:**\n\n```\n{_last_output}\n```", icon="✅")
                    _last_errors = st.session_state.get("_wf_last_errors", "")
                    if _last_errors:
                        st.error(f"```\n{_last_errors}\n```", icon="❌")
                    
                    # Show all processed files from zProcessed directory
                    if os.path.isdir(_processed_dir):
                        _all_processed_files = []
                        for _date_dir in sorted(os.scandir(_processed_dir), key=lambda d: d.name, reverse=True):
                            if _date_dir.is_dir():
                                for _pf in sorted(os.scandir(_date_dir.path), key=lambda f: f.name):
                                    if _pf.is_file():
                                        _all_processed_files.append((_date_dir.name, _pf.name))
                        if _all_processed_files:
                            st.markdown(f"**✅ Processed** ({len(_all_processed_files)} files):")
                            for _date_label, _fname in _all_processed_files:
                                st.markdown(f"&nbsp;&nbsp;&nbsp;&nbsp;📄 `{_fname}` &nbsp;·&nbsp; {_date_label}")
                    
                    # Hidden rerun trigger (inside collapsed expander = naturally hidden)
                    _wf_rerun_clicked = st.button("⟳", key="_wf_rerun_trigger", type="secondary")
                
                # Persistently load processed PDFs into PROCESSED RESULT section
                # Only runs if processing happened this session (not on initial page load)
                import datetime as _wf_dt2
                _processed_dir2 = os.path.join(_resolved_watch_path, "zProcessed")
                _wf_has_run_this_session = bool(st.session_state.get("_wf_last_output"))
                if os.path.isdir(_processed_dir2) and _wf_has_run_this_session and not get_skill_state("last_processed_files"):
                    _wf_all_pdfs = []
                    # Scan all date directories (newest first)
                    try:
                        _date_dirs = sorted(os.scandir(_processed_dir2), key=lambda d: d.name, reverse=True)
                    except OSError:
                        _date_dirs = []
                    for _dd in _date_dirs:
                        if _dd.is_dir():
                            try:
                                for _pe in sorted(os.scandir(_dd.path), key=lambda e: e.name):
                                    if _pe.is_file() and _pe.name.lower().endswith(".pdf"):
                                        try:
                                            with open(_pe.path, "rb") as _rpf:
                                                _rpf_bytes = _rpf.read()
                                            _wf_all_pdfs.append({
                                                "name": _pe.name,
                                                "bytes": _rpf_bytes,
                                                "transcript": f"✅ Converted: {_pe.name} ({_dd.name})",
                                                "content_preview": "",
                                            })
                                        except (OSError, IOError):
                                            pass
                            except OSError:
                                pass
                    if _wf_all_pdfs:
                        set_skill_state("last_processed_files", _wf_all_pdfs)
                        _wf_summary = f"Watch folder: {len(_wf_all_pdfs)} file(s) converted to PDF"
                        set_skill_state("last_output", _wf_summary)
                        set_skill_state("file_index", 0)
                
                # JS timer — finds button inside expander DOM and clicks it on interval
                st.components.v1.html(
                    f"""
                    <script>
                    (function() {{
                        const doc = window.parent.document;
                        function findRerunBtn() {{
                            const btns = doc.querySelectorAll('button[kind="secondary"]');
                            for (const b of btns) {{
                                if (b.textContent.trim() === '⟳') return b;
                            }}
                            return null;
                        }}
                        if (window.parent._watchFolderTimer) {{
                            clearInterval(window.parent._watchFolderTimer);
                        }}
                        window.parent._watchFolderTimer = setInterval(function() {{
                            const btn = findRerunBtn();
                            if (btn) btn.click();
                        }}, {_watch_interval_ms});
                    }})();
                    </script>
                    """,
                    height=0
                )
        else:
            st.caption("Enter a folder path above to enable automatic file processing.")
    else:
        # Clear timer when toggle is off
        st.components.v1.html(
            """
            <script>
            if (window._watchFolderTimer) {
                clearInterval(window._watchFolderTimer);
                window._watchFolderTimer = null;
            }
            </script>
            """,
            height=0
        )
    st.markdown("---")

col_head, col_toggle = st.columns([1, 1])
if is_audio_skill:
    with col_head:
        st.subheader("Upload Audio Files", anchor=False)
    uploader_label = "Upload Audio Files"
    accepted_types = [
        "mp3", "wav", "m4a", "aac", "ogg", "flac", "webm",
        "aiff", "aif", "wma", "oga", "opus", "3gp",
        "mp4", "mov", "avi", "mkv"
    ]
elif is_tts_skill:
    with col_head:
        st.subheader("Upload Text Files for Narration", anchor=False)
    uploader_label = "Upload Text Files"
    st.caption("ElevenLabs cleanly auto-extracts text from Plain Text, Markdown, RTF, DOC, and DOCX files for narration.")
    accepted_types = [
        "txt", "md", "rtf", "doc", "docx", "csv", "json", "py", "sh", "yaml", "yml", "ini"
    ]
else:
    with col_head:
        st.subheader("Upload Document Files", anchor=False)
    uploader_label = "Upload Document Files"
    accepted_types = [
        "txt", "md", "docx", "doc", "csv", "json", "rtf", "py", "sh", "yaml", "yml"
    ]

safety_net_on = False
if is_audio_skill or is_tts_skill:
    with col_toggle:
        st.markdown("<div style='margin-top: 5px;'></div>", unsafe_allow_html=True)
        safety_net_on = st.toggle("SAVE TO GOOGLE ACCT", value=False, help="Saves to Google Drive & Sheet")
        
        # Style the toggle to be single-line, right-aligned
        st.markdown("""
            <style>
            /* Force the toggle label to be on a single line and right align the whole component */
            div[data-testid="column"]:nth-child(2) {
                display: flex;
                justify-content: flex-end;
                align-items: center;
            }
            div[data-testid="column"]:nth-child(2) div[data-testid="stToggle"] {
                white-space: nowrap;
                width: max-content;
                float: right;
            }
            </style>
        """, unsafe_allow_html=True)
        
        # Toggle green color is handled by CSS hue-rotate filter in style.css (section 19)

# --- Duplicate Checking Logic ---
def check_new_uploads_for_duplicates(file_list):
    """
    Checks incoming files against processed files.
    Triggers an error if duplicates exist and RETURNS a filtered list of only NEW files.
    """
    import collections
    if not file_list:
        set_skill_state("prev_file_counts_dict", collections.Counter())
        return file_list
        
    processed_raw = get_skill_state("processed_files", set())
    processed: set[str] = processed_raw if isinstance(processed_raw, set) else set()
    
    # Count ALL files currently in the widget
    current_files = [f.name + str(f.size) for f in file_list]
    curr_counts: dict[str, int] = dict(collections.Counter(current_files))
    
    # Prime state on cold start
    ns_key = f"{st.session_state.selected_skill_id}_prev_file_counts_dict"
    if ns_key not in st.session_state:
        # Treat as empty previous state so all dragged files are considered "new" and will trip the duplicate check if already processed
        prev_counts: dict[str, int] = {}
    else:
        prev_counts_raw = get_skill_state("prev_file_counts_dict", {})
        prev_counts = dict(prev_counts_raw) if isinstance(prev_counts_raw, dict) else {}  # type: ignore[arg-type]
    
    error_triggered = False
    clean_list = []
    
    for f in file_list:
        file_id = f.name + str(f.size)
        # If this exact file ID is in our totally processed set, it's a duplicate
        if file_id in processed: # type: ignore
            # We also check if the user *just* added it (count increased), which triggers the visual error
            # If they just hit refresh and the count is the same, we silently remove it without screaming
            if curr_counts.get(file_id, 0) > prev_counts.get(file_id, 0): # type: ignore
                error_triggered = True
        else:
            # Not processed yet, keep it!
            clean_list.append(f)
            
    if error_triggered:
        trigger_duplicate_error()
        # We simply drop the file from the processing queue, but don't stop the script.
        # This prevents locking the UI out of processing newly dragged valid files.
    # Update tracking state with what the UI actually holds right now
    set_skill_state("prev_file_counts_dict", curr_counts)
    
    return clean_list

# File uploader OUTSIDE the form so uploads trigger immediately
skill_args: dict[str, Any] = {}

# --- Specialized Skill Inputs (Part 1: Above Uploader) ---
if selected_skill and selected_skill.get("basename") in ["AI-LLM-Speech2Text", "AI-LLM-KIE-ElevenLabs-Speech2Text", "AI-LLM-Text2Speech", "AI-LLM-KIE-ElevenLabs-Text2Speech"]:
    def copy_folder_to_sheet():
        st.session_state["google_sheet_input"] = st.session_state.get("drive_folder_input", "")

    if safety_net_on:
        # Create a hidden marker div to anchor our CSS sibling selectors
        st.markdown('<div id="safety-net-marker" style="display:none;"></div>', unsafe_allow_html=True)
        st.markdown(
            """
            <style>
            @keyframes softFadeIn {
                0% { opacity: 0; transform: translateY(-8px); }
                100% { opacity: 1; transform: translateY(0); }
            }
            /* Target the columns container and the email text_input that immediately follow the marker */
            /* Streamlit wraps markdown in an iframe-like div, so we target the parent's siblings */
            div[data-testid="stMarkdownContainer"]:has(#safety-net-marker) {
                display: none;
            }
            div[data-testid="stVerticalBlock"] > div:has(> div > div > div > div#safety-net-marker) ~ div {
                animation: softFadeIn 1.5s cubic-bezier(0.2, 0.8, 0.2, 1) forwards;
            }
            /* Specifically target the next 2 blocks (columns and email input) just in case */
            div[data-testid="stVerticalBlock"] > div:has(> div > div > div > div#safety-net-marker) + div,
            div[data-testid="stVerticalBlock"] > div:has(> div > div > div > div#safety-net-marker) + div + div {
                animation: softFadeIn 1.5s cubic-bezier(0.2, 0.8, 0.2, 1) forwards;
            }
            </style>
            """,
            unsafe_allow_html=True
        )
        
        col1, col_btn, col2 = st.columns([3, 1, 3])
        with col1:
            skill_args["drive_folder"] = st.text_input(
                "Google Drive Folder (Optional):", 
                placeholder="e.g. AI-Audio/Podcasts",
                key="drive_folder_input"
            )
        with col_btn:
            st.write("") # Spacer
            st.button(
                "Copy ➜", 
                on_click=copy_folder_to_sheet, 
                help="Copy Folder Name to Sheet Name",
                use_container_width=True
            )
            
            # Streamlit doesn't support custom button colors natively without 'type'.
            # We target the button specifically within this column to make it yellow and push it down to align with inputs.
            st.markdown(
                """
                <style>
                div[data-testid="column"]:nth-of-type(2) {
                    display: flex;
                    align-items: flex-end; /* push to bottom of row to match input boxes */
                    padding-bottom: 2px;
                }
                div[data-testid="column"]:nth-of-type(2) .stButton > button {
                    background-color: #ffd700 !important;
                    color: #111 !important;
                    font-weight: 700 !important;
                    border: none !important;
                    transition: transform 0.2s !important;
                    margin-top: 28px !important; /* height of the label above inputs */
                }
                div[data-testid="column"]:nth-of-type(2) .stButton > button:hover {
                    background-color: #ffe44d !important;
                    color: #000 !important;
                    border: none !important;
                    transform: scale(1.02) !important;
                }
                </style>
                """,
                unsafe_allow_html=True
            )
        with col2:
            skill_args["google_sheet"] = st.text_input(
                "Google Sheet Name (Optional):", 
                placeholder="e.g. Transcription Database",
                key="google_sheet_input"
            )
        
        default_email = st.session_state.get("GCP_USER_EMAIL", os.environ.get("GCP_USER_EMAIL", ""))
        skill_args["share_with"] = st.text_input("User Email for Auto-Sharing (Optional):", 
                                                value=default_email,
                                                type="password",
                                                placeholder="e.g. user@gmail.com")
    # Google Drive/Sheet direct-link badges are shown in the PROCESSED RESULT section only
    # (see show_result_popup below)
if selected_skill and selected_skill.get("basename") in ["Data-GoogleSheet", "Data-CustomGoogleSheet"]:
    uploaded_files = []
    # Optionally display a nice instruction block instead of the uploader
    st.info("📊 **Google Sheet Generator**\nConfigure your desired spreadsheet below.")
else:
    # Restrict to CSV only if the CSV-to-Sheet skill is selected
    allowed_types = ["csv"] if selected_skill and selected_skill.get("basename") == "Data-CSV2GoogleSheet" else accepted_types
    
    uploaded_files = st.file_uploader(
        uploader_label,
        accept_multiple_files=True,
        label_visibility="collapsed",
        type=allowed_types
    )



# --- PDF Restriction Safeguard ---
if uploaded_files:
    pdf_files = [f for f in uploaded_files if f.name.lower().endswith(".pdf")]
    if pdf_files:
        st.warning("⚠️ **PDF files are not allowed for upload.** These files have been removed from your selection.")
        uploaded_files = [f for f in uploaded_files if not f.name.lower().endswith(".pdf")]

# Keybind ENTER to the Browse files button
st.components.v1.html(
    """
    <script>
    const doc = window.parent.document;
    if (!doc._enterKeyBound) {
        doc._enterKeyBound = true;
        doc.addEventListener('keydown', function(e) {
            if (e.key === 'Enter' && e.target.tagName !== 'INPUT' && e.target.tagName !== 'TEXTAREA') {
                // Don't trigger Browse if a result is being displayed or a button has focus
                const resultHeader = Array.from(doc.querySelectorAll('h2')).find(h => h.innerText.includes('PROCESSED RESULT'));
                if (resultHeader || doc.querySelector('button.kb-focus')) return;
                // Try multiple selectors to find the Browse button
                const browseBtn = doc.querySelector('[data-testid="stFileUploaderDropzone"] button')
                    || doc.querySelector('[data-testid="baseButton-secondary"]')
                    || doc.querySelector('section[data-testid="stFileUploader"] button');
                if (browseBtn) {
                    browseBtn.click();
                    e.preventDefault();
                }
            }
        });
    }
    </script>
    """,
    height=0
)
if uploaded_files and not is_audio_skill and not is_tts_skill:
    # Server-side filter: remove any .pdf files that slipped through (e.g. from prior session)
    uploaded_files = [f for f in uploaded_files if not f.name.lower().endswith('.pdf')]

# ALWAYS check for duplicates to synchronize state when the user clears the widget
uploaded_files = check_new_uploads_for_duplicates(uploaded_files if uploaded_files else [])

# --- Merge in files from the ADD MORE uploader in the result popup ---
_addmore_pending = st.session_state.pop("_addmore_pending", None)
if _addmore_pending:
    # Combine with any existing uploads, avoiding duplicates by name+size
    existing_ids = set((f.name + str(f.size)) for f in (uploaded_files or []))
    # Also exclude files already processed in this session
    _proc_set = get_skill_state("processed_files", set())
    if isinstance(_proc_set, set):
        existing_ids.update(_proc_set)
    for af in _addmore_pending:
        if (af.name + str(af.size)) not in existing_ids:
            if uploaded_files is None:
                uploaded_files = []
            uploaded_files.append(af)
            existing_ids.add(af.name + str(af.size))
        
if uploaded_files:
    file_names = ", ".join([f.name for f in uploaded_files])
    set_skill_state("_upload_status_msg", f"📎 {len(uploaded_files)} file(s) uploaded: **{file_names}**")

# Detect if a NEW file was just uploaded (auto-run trigger)
# Auto-run for ALL files over riding the previous logic
current_upload_id = None
if uploaded_files:
    current_upload_id = "|".join(sorted(f.name + str(f.size) for f in uploaded_files))

auto_run = False
if current_upload_id and current_upload_id != get_skill_state("prev_upload_id"):
    set_skill_state("prev_upload_id", current_upload_id)
    
    # Filter out files that have already been processed in this session
    processed = get_skill_state("processed_files", set())
    new_files = [f for f in uploaded_files if (f.name + str(f.size)) not in processed]

    if new_files:
        auto_run = True

manual_run_clicked = False

# --- URL Input for Specific Skills ---
url_input = ""
html_capture = False
if url_input:
    st.info(f"Targeting URL: `{url_input}`")

# --- Specialized Skill Inputs (Part 2: Below Uploader) ---
if selected_skill and selected_skill.get("basename") == "AI-LLM-ImageGenerate":
    skill_args["prompt"] = st.text_area("Image Prompt:", placeholder="A futuristic city with neon lights...", height=100)
    col1, col2 = st.columns(2)
    with col1:
        skill_args["count"] = st.number_input("Number of Images:", min_value=1, max_value=4, value=1)
    with col2:
        skill_args["provider"] = st.selectbox("Provider:", ["gemini", "openai"])
elif selected_skill and selected_skill.get("basename") == "AI-LLM-EmbedText":
    mode = st.radio("Mode:", ["Single Text", "Compare Two Texts"], horizontal=True)
    if mode == "Single Text":
        skill_args["text"] = st.text_area("Text to Embed:", height=150)
    else:
        skill_args["compare"] = [
            st.text_input("Text A:"),
            st.text_input("Text B:")
        ]
elif selected_skill and selected_skill.get("basename") == "AI-LLM-RAGQuery":
    skill_args["query"] = st.text_input("Your Question:", placeholder="What does this document say about...?")
    skill_args["index"] = st.checkbox("Re-index Documents", value=True)
elif selected_skill and selected_skill.get("basename") == "AI-LLM-TranslateText":
    skill_args["to"] = st.text_input("Target Language:", value="Spanish")
    # If no file, show text area
    if not uploaded_files:
        skill_args["text"] = st.text_area("Text to Translate:", height=150)
elif selected_skill and selected_skill.get("basename") in ["Data-GoogleSheet", "Data-CustomGoogleSheet"]:
    # Proactive check for credentials.json
    creds_path = os.path.join(str(selected_skill["dir"]), "credentials.json")
    if not os.path.exists(creds_path):
        pass
        
    st.markdown('<div class="google-sheet-input-wrapper">', unsafe_allow_html=True)
    skill_args["title"] = st.text_input("Google Sheet Title:", placeholder="e.g. Q3 Sales Tracking")
    skill_args["fields"] = st.text_input("Column Headers (comma separated):", placeholder="e.g. First Name, Last Name, Email, Phone")
    st.markdown('</div>', unsafe_allow_html=True)
    
    st.markdown("<br>", unsafe_allow_html=True)
    col1, col2, col3 = st.columns([1, 2, 1])
    with col2:
        enter_clicked = st.button("✨ GENERATE GOOGLE SHEET", use_container_width=True, type="primary")
        if enter_clicked:
            skill_args["_force_run"] = "True" # type: ignore
            
        # JS to Bind Return key (Enter) to the button specifically for these inputs
        st.markdown(
            """
            <script>
            {
                const doc = window.parent.document;
                const bindEnter = () => {
                    const wrapper = doc.querySelector('.google-sheet-input-wrapper');
                    if (wrapper) {
                        const inputs = wrapper.querySelectorAll('input');
                        inputs.forEach(input => {
                            if (!input._gsBound) {
                                input._gsBound = true;
                                input.addEventListener('keydown', function(ev) {
                                    if (ev.key === 'Enter') {
                                        ev.preventDefault();
                                        const buttons = Array.from(doc.querySelectorAll('button'));
                                        const genBtn = buttons.find(b => b.innerText.includes('GENERATE GOOGLE SHEET'));
                                        if (genBtn) genBtn.click();
                                    }
                                });
                            }
                        });
                    }
                };
                
                // Try binding immediately and also a few times after UI refreshes
                bindEnter();
                setTimeout(bindEnter, 500);
                setTimeout(bindEnter, 1000);
            }
            </script>
            """,
            unsafe_allow_html=True
        )

# --- Manual Text Input Fallback ---
manual_text = ""
# Check for any values in skill_args EXCEPT drive_folder
has_other_skill_args = any(v for k, v in skill_args.items() if k != "drive_folder")
if not uploaded_files and not url_input and not has_other_skill_args:
    if selected_skill["basename"] not in [
        "AI-LLM-ImageGenerate", 
        "AI-LLM-EmbedText", 
        "AI-LLM-RAGQuery", 
        "AI-LLM-TranslateText",
        "AI-LLM-Speech2Text",
        "AI-LLM-KIE-ElevenLabs-Speech2Text",
        "Data-GoogleSheet",
        "Data-CustomGoogleSheet",
        "Data-CSV2GoogleSheet"
    ]:
        if is_tts_skill:
            with st.form("manual_tts_input_form", clear_on_submit=True):
                st.markdown('<div class="tts-manual-input-wrapper">', unsafe_allow_html=True)
                manual_text = st.text_area("Manual Text Input (Optional):", placeholder="Paste text here to process instead of uploading a file...", height=200)
                st.markdown('</div>', unsafe_allow_html=True)
                
                # JS to Bind Enter key (without Shift) to the form submission
                st.markdown("""
                    <script>
                    const doc = window.parent.document;
                    const wrapper = doc.querySelector('.tts-manual-input-wrapper');
                    if (wrapper) {
                        const textarea = wrapper.querySelector('textarea');
                        if (textarea) {
                            textarea.addEventListener('keydown', function(e) {
                                if (e.key === 'Enter' && !e.shiftKey) {
                                    e.preventDefault();
                                    const submitBtn = wrapper.closest('form').querySelector('button[kind="formSubmit"]');
                                    if (submitBtn) submitBtn.click();
                                }
                            });
                        }
                    }
                    </script>
                """, unsafe_allow_html=True)
                
                st.markdown('<div class="tts-enter-marker"></div>', unsafe_allow_html=True)
                enter_clicked = st.form_submit_button("ENTER", use_container_width=False)
        else:
            manual_text = st.text_area("Manual Text Input (Optional):", placeholder="Paste text here to process instead of uploading a file...", height=200)
            enter_clicked = False
    else:
        enter_clicked = False

# Handle Execution
has_special_input = False
if selected_skill and selected_skill.get("basename") == "AI-LLM-ImageGenerate" and skill_args.get("prompt"):
    has_special_input = True
elif selected_skill and selected_skill.get("basename") == "AI-LLM-EmbedText" and (skill_args.get("text") or (skill_args.get("compare") and all(skill_args["compare"]))):
    has_special_input = True
elif selected_skill and selected_skill.get("basename") == "AI-LLM-RAGQuery" and skill_args.get("query"):
    has_special_input = True
elif selected_skill and selected_skill.get("basename") == "AI-LLM-TranslateText" and (skill_args.get("to") and (uploaded_files or skill_args.get("text"))):
    has_special_input = True
elif selected_skill and selected_skill.get("basename") in ["Data-GoogleSheet", "Data-CustomGoogleSheet"] and skill_args.get("title") and skill_args.get("fields") and skill_args.get("_force_run") == "True":
    has_special_input = True

# Prevent infinite auto-run loops when clearing the results popup
auto_manual = False
if not is_tts_skill and manual_text != "":
    if manual_text != get_skill_state("prev_manual_text"):
        set_skill_state("prev_manual_text", manual_text)
        auto_manual = True

auto_special = False
if has_special_input:
    curr_spec = str(skill_args)
    if curr_spec != get_skill_state("prev_spec_args"):
        set_skill_state("prev_spec_args", curr_spec)
        auto_special = True

auto_url = False
if url_input != "":
    if url_input != get_skill_state("prev_url_input"):
        set_skill_state("prev_url_input", url_input)
        auto_url = True

should_run = auto_run or manual_run_clicked or auto_url or auto_special or (manual_text != "" and enter_clicked) or auto_manual

if should_run:
    args_input = ""
    
    # Construct args from specialized inputs
    if selected_skill and selected_skill.get("basename") == "AI-LLM-ImageGenerate":
        args_input = f"--prompt {shlex.quote(str(skill_args['prompt']))} --count {skill_args['count']} --provider {str(skill_args['provider'])}"
    elif selected_skill and selected_skill.get("basename") == "AI-LLM-EmbedText":
        if skill_args.get("text"):
            args_input = f"--text {shlex.quote(str(skill_args['text']))}"
        elif skill_args.get("compare"):
            args_input = f"--compare {shlex.quote(str(skill_args['compare'][0]))} {shlex.quote(str(skill_args['compare'][1]))}"
    elif selected_skill and selected_skill.get("basename") == "AI-LLM-RAGQuery":
        args_input = f"--query {shlex.quote(str(skill_args['query']))}"
        if skill_args.get("index"):
            args_input += " --index"
        if uploaded_files:
            # We'll handle files below
            pass
    elif selected_skill and selected_skill.get("basename") == "AI-LLM-TranslateText":
        args_input = f"--to {shlex.quote(str(skill_args['to']))}"
        if skill_args.get("text"):
            # We'll save this to a temp file below
            pass
    elif selected_skill and selected_skill.get("basename") in ["Data-GoogleSheet", "Data-CustomGoogleSheet"]:
        fields = [shlex.quote(f.strip()) for f in str(skill_args.get("fields", "")).split(",") if f.strip()]
        fields_str = " ".join(fields)
        args_input = f"--title {shlex.quote(str(skill_args.get('title', '')))} --fields {fields_str}"
    elif selected_skill and selected_skill.get("basename") == "Data-CSV2GoogleSheet":
        args_input = "--file {FILE_1}"
    
    if url_input:
        args_input = f"--url \"{url_input}\""
        if html_capture:
            args_input += " --capture"
            
    # --- Safety Net Fallback Auto-Injection ---
    if safety_net_on:
        valid_skills_for_safety = ["AI-LLM-Speech2Text", "AI-LLM-KIE-ElevenLabs-Speech2Text", "AI-LLM-Text2Speech", "AI-LLM-KIE-ElevenLabs-Text2Speech"]
        if selected_skill and selected_skill.get("basename") in valid_skills_for_safety:
            need_fallback = not skill_args.get("drive_folder") or not skill_args.get("google_sheet")
            if need_fallback:
                # Initialize or get increment counter
                idx = st.session_state.get("_unnamed_counter", 1)
                default_name = f"UnNamed-{idx}"
                
                if not skill_args.get("drive_folder"):
                    skill_args["drive_folder"] = default_name
                if not skill_args.get("google_sheet"):
                    skill_args["google_sheet"] = default_name
                    
                # Increment for next time
                st.session_state["_unnamed_counter"] = idx + 1
    
    # Handle uploaded files by saving them to a temporary directory so the script can read them
    temp_dir = None
    file_paths = []
    
    if uploaded_files or manual_text or skill_args.get("text"):
        # Filter uploaded_files to only process NEW ones
        processed = get_skill_state("processed_files", set())
        
        # Combine uploaded files and manual text
        files_to_process_objs = [uf for uf in uploaded_files if (uf.name + str(uf.size)) not in processed]
        
        # Deduplicate by name+size (Streamlit uploader can accumulate duplicate entries)
        _seen_ids: set[str] = set()
        _deduped: list = []
        for uf in files_to_process_objs:
            _fid = uf.name + str(uf.size)
            if _fid not in _seen_ids:
                _seen_ids.add(_fid)
                _deduped.append(uf)
        files_to_process_objs = _deduped
        
        if not files_to_process_objs and not url_input and not manual_text and not skill_args.get("text") and not has_special_input:
            should_run = False # Cancel execution if no new files, no URL, no manual text
        
        else:
            temp_dir = tempfile.mkdtemp()
            
            # Handle Manual Text Input
            if (manual_text or skill_args.get("text")) and not files_to_process_objs:
                text_content = manual_text or skill_args.get("text")
                
                # Use a counter to ensure unique filenames for manual entries
                if "manual_tts_counter" not in st.session_state:
                    st.session_state["manual_tts_counter"] = 1
                else:
                    st.session_state["manual_tts_counter"] += 1
                
                counter = st.session_state["manual_tts_counter"]
                text_file_path = os.path.join(temp_dir, f"input_text_{counter}.txt")
                
                with open(text_file_path, "w", encoding="utf-8") as f:
                    f.write(str(text_content))
                file_paths.append(text_file_path)
                
                # For RAGQuery, --docs is required
                if selected_skill_id == "AI-LLM-RAGQuery":
                    args_input += f" --docs {text_file_path}"
                # For Translate/General, if args empty, use as --input
                elif not args_input.strip() or "--to" in args_input:
                     args_input += f" --input {text_file_path}"
            
            # Handle Uploaded Files
            elif files_to_process_objs:
                if not args_input.strip() or ("--to" in args_input and "--input" not in args_input):
                    # Only auto-inject if it's not already specialized except for Translate which needs --input
                    if selected_skill_id == "AI-LLM-RAGQuery":
                         args_input += " --docs {FILE_1}"
                    else:
                         args_input += " --input {FILE_1}"
                
                # Preserve the base args so we don't lose the {FILE_1} placeholder when looping
                base_args_input = str(args_input)
                
                for i, uf in enumerate(files_to_process_objs):
                    file_path = os.path.join(temp_dir, uf.name)
                    with open(file_path, "wb") as f:
                        f.write(uf.getbuffer())
                    file_paths.append(file_path)
                    
                    # We NO LONGER mark this as processed here. We wait until successful execution.
                    # processed.add(uf.name + str(uf.size)) # type: ignore
                    # set_skill_state("processed_files", processed)
                    
                    # Replace {FILE_X} placeholder in the args for THIS SPECIFIC file (for single-run fallback)
                    # We only do this for the *first* file because single execution assumes args_input is fully resolved.
                    # The batch loop relies on base_args_input later.
                    if i == 0:
                        args_input = base_args_input.replace("{FILE_1}", file_path)
    
    # Parse the arguments string into a list safely avoiding simple split() issues with quotes
    try:
        parsed_args = shlex.split(args_input)
    except Exception as e:
        st.error(f"Error parsing arguments: {e}")
        st.stop()
        
    python_cmd: str = str(get_python_cmd())
    script_path: str = str(selected_skill["script"])
    command: list[str] = [python_cmd, script_path] + parsed_args
    
    cwd: str = str(selected_skill["dir"])
    
    # Use manual status instead of with st.spinner so we can clear it before st.stop()
    
    # NOTE: We no longer reset last_processed_files here — clips accumulate
    # across batches within the session. The file_index is set after processing
    # to point at the start of the newest batch (see line ~2231).
    
    proc_overlay = trigger_processing_overlay()
    output_expander = st.expander("📄 Data", expanded=False)
    with output_expander:
        main_spinner = st.empty()
        main_spinner.info("⏳ PROCESSING...")
        try:
            # We stream the output to a preformatted block
            # For simplicity in Streamlit, we capture combined output
            # Inject API keys from session state (auto-fetched from GCP at login)
            run_env = os.environ.copy()
            if st.session_state.get("GEMINI_API_KEY"):
                run_env["GEMINI_API_KEY"] = st.session_state["GEMINI_API_KEY"]
            if st.session_state.get("ELEVENLABS_API_KEY"):
                run_env["ELEVENLABS_API_KEY"] = st.session_state["ELEVENLABS_API_KEY"]
            if st.session_state.get("KIE_API_KEY"):
                run_env["KIE_API_KEY"] = st.session_state["KIE_API_KEY"]
            
            if is_audio_skill:
                new_files = process_uploaded_files(file_paths, selected_skill, run_env, 
                                                   base_args_input=locals().get("base_args_input", args_input),
                                                   drive_folder=skill_args.get("drive_folder", ""), 
                                                   google_sheet=skill_args.get("google_sheet", ""),
                                                   share_with=skill_args.get("share_with", ""),
                                                   proc_overlay=proc_overlay, main_spinner=main_spinner)
                existing = get_skill_state("last_processed_files", [])
                existing.extend(new_files)
                set_skill_state("last_processed_files", existing)
                
                if new_files:
                    # Successfully processed these files, safe to add them to the processed duplicate tracker
                    processed = get_skill_state("processed_files", set())
                    for f in new_files:
                        original_name = str(f.get("original_name", ""))
                        if original_name:
                            # Match it back to the uploaded_file to get the exact size footprint
                            for uf in uploaded_files:
                                if uf.name == original_name:
                                    processed.add(uf.name + str(uf.size)) # type: ignore
                                    break
                    set_skill_state("processed_files", processed)
                    
                    set_skill_state("file_index", max(0, len(existing) - len(new_files)))
                    set_skill_state("last_output", new_files[0]["transcript"])
                    set_skill_state("auto_open_result", True)
                    _msg = f"✅ Successfully processed {len(file_paths)} file(s)"
                    st.success(_msg)
                    set_skill_state("_last_success_msg", _msg)
                    # Store combined processing output for persistent Data expander
                    _all_output = []
                    for nf in new_files:
                        _all_output.append(str(nf.get("transcript", "")))
                    set_skill_state("_last_data_output", "\n\n".join(_all_output))
            elif is_tts_skill and file_paths:
                new_files = process_tts_files(file_paths, selected_skill, run_env, 
                                              base_args_input=locals().get("base_args_input", args_input),
                                              drive_folder=skill_args.get("drive_folder", ""), 
                                              google_sheet=skill_args.get("google_sheet", ""),
                                              share_with=skill_args.get("share_with", ""),
                                              proc_overlay=proc_overlay, main_spinner=main_spinner)
                existing = get_skill_state("last_processed_files", [])
                existing.extend(new_files)
                set_skill_state("last_processed_files", existing)
                if new_files:
                    set_skill_state("file_index", max(0, len(existing) - len(new_files)))
                    set_skill_state("last_output", new_files[0]["transcript"])
                    set_skill_state("auto_open_result", True)
                    _msg = f"✅ Successfully converted {len(file_paths)} document(s) to audio"
                    st.success(_msg)
                    set_skill_state("_last_success_msg", _msg)
            else:
                # Standard execution PATH
                all_execution_results = []
                
                # Check if we should loop (multiple files and {FILE_1} placeholder exists in base_args)
                # This ensures skills like Convtr-PlainTxt2PDF process ALL uploads.
                # If base_args_input wasn't created (e.g., text area input), it falls back to single execution.
                # --- NEW: Streaming Execution Engine ---
                def run_and_stream(cmd, cwd, env):
                    """Helper to run a command and stream its output to the UI."""
                    process = subprocess.Popen(
                        cmd, cwd=cwd, env=env,
                        stdout=subprocess.PIPE, stderr=subprocess.STDOUT,
                        text=True, bufsize=1, universal_newlines=True
                    )
                    
                    full_output: list[str] = []
                    output_placeholder = st.empty()
                    
                    try:
                        # Read output line by line as it arrives
                        stdout_pipe = process.stdout
                        if stdout_pipe:
                            for line in iter(stdout_pipe.readline, ""):
                                full_output.append(str(line))
                                # Update the UI with the latest 20 lines to keep it snappy
                                output_placeholder.code("".join(full_output[-20:])) # type: ignore
                        
                        returncode = process.wait()
                    finally:
                        if process.poll() is None:
                            process.terminate()
                            try:
                                process.wait(timeout=2)
                            except subprocess.TimeoutExpired:
                                process.kill()
                                
                    final_text = "".join(full_output)
                    
                    # Create a compatibility object for the existing result loop
                    class ProcessResult:
                        def __init__(self, stdout, stderr, returncode):
                            self.stdout = stdout
                            self.stderr = stderr
                            self.returncode = returncode
                            
                    return ProcessResult(final_text, final_text if returncode != 0 else "", returncode)

                base_loop_args = locals().get('base_args_input', args_input)
                should_loop = len(file_paths) > 1 and "{FILE_1}" in str(base_loop_args)
                
                if should_loop:
                    for i, fp in enumerate(file_paths):
                        local_args = str(base_loop_args).replace("{FILE_1}", str(fp))
                        try:
                            local_parsed = shlex.split(local_args)
                            local_cmd = [python_cmd, script_path] + local_parsed
                            main_spinner.info(f"⏳ PROCESSING file {i+1}/{len(file_paths)}: `{os.path.basename(fp)}`...")
                            res = run_and_stream(local_cmd, cwd, run_env)
                            all_execution_results.append((fp, res))
                        except Exception as e:
                            st.error(f"Error processing {fp}: {e}")
                else:
                    # Single execution (Google Sheets, Image Gen, etc.)
                    res = run_and_stream(command, cwd, run_env)
                    all_execution_results.append((file_paths[0] if file_paths else None, res))

                # Now aggregate all results into the playlist
                last_processed_files = []
                
                for fp, result in all_execution_results:
                    if result.returncode == 0:
                        output_text = result.stdout.strip()
                        
                        # Detect generated files for this specific execution
                        res_bytes = None
                        res_name = None
                        saved_paths = []
                        for line in output_text.splitlines():
                            if "Saved:" in line:
                                path_str = str(line.split("Saved:")[1].strip())
                                full_path = path_str if os.path.isabs(path_str) else os.path.join(str(cwd), path_str)
                                if os.path.exists(full_path):
                                    saved_paths.append(full_path)
                                    
                        if saved_paths:
                            # Capture the first saved file as the primary result for this entry
                            res_name = os.path.basename(saved_paths[0])
                            with open(saved_paths[0], "rb") as f:
                                res_bytes = f.read()

                        # Capture preview for the input file (if it exists)
                        content_preview = ""
                        if fp:
                            ext = os.path.splitext(fp)[1].lower()
                            if ext in [".txt", ".rtf", ".md", ".docx", ".doc", ".rtfd"] or (ext != ".pdf" and os.path.getsize(fp) < 1024 * 1024):
                                content_preview = read_text_file_preview(str(fp))

                        # Clean output (filter success messages)
                        ignore_prefixes = ("Transcribing:", "Saved:", "Fetching:", "Capturing:", "Capturing high-fidelity", "Usage:")
                        lines = [l for l in output_text.splitlines() if not l.startswith(ignore_prefixes)]
                        clean_output = "\n".join(lines).strip()
                        
                        # Extract usage stats
                        usage_line = next((l for l in output_text.splitlines() if l.startswith("Usage:")), None)
                        usage_details = ""
                        if usage_line:
                            usage_details = f"\n\n**Statistics:** {usage_line.split(':', 1)[-1].strip()}"

                        last_processed_files.append({
                            "original_name": os.path.basename(fp) if fp else (res_name or "Result"),
                            "name": os.path.basename(fp) if fp else (res_name or "Result"),
                            "bytes": open(fp, "rb").read() if fp and os.path.exists(fp) else (res_bytes or b""),
                            "transcript": (clean_output + usage_details) if clean_output or usage_details else f"✅ Generated: {res_name}",
                            "content_preview": content_preview,
                            "result_bytes": res_bytes,
                            "result_name": res_name
                        })
                        
                        # Mark this specific file as successfully processed so it trips the duplicate checker next time
                        if fp:
                            processed = get_skill_state("processed_files", set())
                            original_name = os.path.basename(fp)
                            for uf in uploaded_files:
                                if uf.name == original_name:
                                    processed.add(uf.name + str(uf.size)) # type: ignore
                                    break
                            set_skill_state("processed_files", processed)
                    else:
                        # Handle errors (Quota, etc.)
                        if "__ANTIGRAVITY_API_QUOTA_EXCEEDED__" in result.stderr:
                            main_spinner.empty()
                            if proc_overlay:
                                proc_overlay.empty()
                            import re
                            quota_msg = ""
                            
                            if "kie" in str(selected_skill["basename"]).lower():
                                match = re.search(r"Kie\.ai Error: (.*?)(?:\n|$)", result.stderr)
                                if match:
                                    quota_msg = f"\n\n**Details:** {match.group(1).strip()}"
                                st.warning(f"⚠️ **DENIED!** You have reached the maximum usage allowed by your **Kie.ai** balance. Please top up your Kie.ai account to continue.{quota_msg}")
                            else:
                                match = re.search(r"['\"]message['\"]:\s*['\"](.*?)['\"]", result.stderr)
                                if match:
                                    quota_msg = f"\n\n**Usage stats:** {match.group(1)}"
                                st.warning(f"⚠️ **DENIED!** You have reached the maximum usage allowed by your **ElevenLabs** active subscription/plan. Please upgrade your plan or wait for the quota to reset.{quota_msg}")
                            st.stop()
                        else:
                            st.error("Execution Error")
                            st.code(result.stderr)
                            if result.stdout:
                                st.code(result.stdout)

                # Final integration
                set_skill_state("last_processed_files", last_processed_files)
                if last_processed_files:
                    set_skill_state("file_index", 0)
                    set_skill_state("last_output", last_processed_files[0]["transcript"])
                    set_skill_state("auto_open_result", True)
                    _msg = f"✅ Successfully processed {len(all_execution_results)} item(s)"
                    st.success(_msg)
                    set_skill_state("_last_success_msg", _msg)
                    
        except Exception as e:
            st.error(f"❌ Error executing skill: {str(e)}")
        finally:
            main_spinner.empty()
            # Result now renders as overlay — no scroll needed
            st.session_state["_delay_autoplay"] = True
            trigger_complete_overlay(proc_overlay)
            time.sleep(0.5)



# --- Persistent Data Expander (hidden via CSS, kept for potential future use) ---
_stored_upload_msg = get_skill_state("_upload_status_msg", "") if selected_skill else ""
_stored_success_msg = get_skill_state("_last_success_msg", "") if selected_skill else ""
_stored_data_output = get_skill_state("_last_data_output", "") if selected_skill else ""
if (_stored_upload_msg or _stored_success_msg or _stored_data_output) and False:  # Hidden — set True to re-enable
    with st.expander("📄 Data", expanded=False):
        if _stored_upload_msg:
            st.success(_stored_upload_msg)
        if _stored_data_output:
            st.code(_stored_data_output)
        if _stored_success_msg:
            st.success(_stored_success_msg)

# -----------------------------------------------------------------------------
# 5. Result Display (Inline instead of Popup)
# -----------------------------------------------------------------------------
def show_result_popup(text: str):
    load_css() # Ensure styles are applied
    processed_files: list = get_skill_state("last_processed_files", []) # type: ignore
    idx: int = int(get_skill_state("file_index", 0))
    selected_skill_id_raw = st.session_state.get("selected_skill_id", "")
    is_google_sheet = any(name in str(selected_skill_id_raw) for name in ["Data-GoogleSheet", "Data-CustomGoogleSheet", "Data-CSV2GoogleSheet"])
    is_media = False
    is_image = False
    
    # 0. Header & Type Info
    # Hide Streamlit's auto-generated anchor link icon on headers
    st.markdown("<style>h1 a, h2 a, h3 a, h4 a { display: none !important; }</style>", unsafe_allow_html=True)
    if is_google_sheet:
        st.markdown("<h1 class='processed-header' style='text-align:center;'><span style='filter:none;'>📄</span> GOOGLE SHEET GENERATED</h1>", unsafe_allow_html=True)
    else:
        st.markdown("<style>.processed-header a {display:none !important;}</style><h1 class='processed-header' style='text-align:center;'><span style='filter:none;'>📄</span> PROCESSED RESULT</h1>", unsafe_allow_html=True)
    
    # Show Google Drive / Sheets badges if IDs were captured during processing
    _popup_has_folder = bool(st.session_state.get("_google_folder_id"))
    _popup_has_sheet = bool(st.session_state.get("_google_sheet_id"))
    if _popup_has_folder or _popup_has_sheet:
        _badge_html = '<style>@keyframes badgeFadeIn{from{opacity:0;transform:translateY(6px)}to{opacity:1;transform:translateY(0)}}</style>'
        _badge_html += '<div style="display:flex;gap:10px;margin:8px 0 12px 0;animation:badgeFadeIn 1s ease-out;">'
        if _popup_has_folder:
            _drive_url = f"https://drive.google.com/drive/folders/{st.session_state['_google_folder_id']}"
            _badge_html += f'<a href="{_drive_url}" target="_blank" style="display:inline-flex;align-items:center;gap:6px;background:#1a73e8;color:#fff;padding:5px 14px;border-radius:20px;font-size:13px;font-weight:600;text-decoration:none;font-family:sans-serif;">📁 Google Drive</a>'
        if _popup_has_sheet:
            _sheet_url = f"https://docs.google.com/spreadsheets/d/{st.session_state['_google_sheet_id']}"
            _badge_html += f'<a href="{_sheet_url}" target="_blank" style="display:inline-flex;align-items:center;gap:6px;background:#0f9d58;color:#fff;padding:5px 14px;border-radius:20px;font-size:13px;font-weight:600;text-decoration:none;font-family:sans-serif;">📊 Google Sheet</a>'
        _badge_html += '</div>'
        st.markdown(_badge_html, unsafe_allow_html=True)

    if processed_files:
        current_file: dict = processed_files[min(idx, len(processed_files)-1)] # type: ignore
        import mimetypes
        file_name = current_file.get("name") or "Result"
        mime_type, _ = mimetypes.guess_type(file_name)
        is_media = mime_type and (mime_type.startswith("audio/") or mime_type.startswith("video/"))
    
    # Identify which text to display (specific clip transcript or the general output)
    display_text = text
    if processed_files:
        current_idx = get_skill_state("file_index", 0)
        if current_idx < len(processed_files):
            current_file = processed_files[current_idx]
            file_transcript = current_file.get("transcript", "")
            content_preview = current_file.get("content_preview", "")
            
            # Identify if the transcript or global text are just status/success messages
            is_file_success = any(x in file_transcript for x in ["Successfully", "Generated:", "✅"]) or not file_transcript.strip()
            is_global_success = any(x in text for x in ["Successfully", "Generated:", "✅"])
            
            if not is_google_sheet:
                if content_preview.strip():
                    # If we have a preview, favor it over any success/batch messages
                    if is_file_success or is_global_success:
                        display_text = content_preview
                    else:
                        display_text = file_transcript or content_preview
                else:
                    # If no preview, only show transcripts if they aren't just success messages
                    # This prevents "✅ Generated: ..." from appearing when navigating between files
                    ft_clean = file_transcript if not any(x in file_transcript for x in ["Successfully", "Generated:", "✅"]) else ""
                    txt_clean = text if not any(x in text for x in ["Successfully", "Generated:", "✅"]) else ""
                    display_text = ft_clean or txt_clean
            else:
                # For Google Sheets, we explicitly NEED the success messages to parse URL/Title
                display_text = file_transcript or text
            
    if processed_files:
        current_file: dict = processed_files[idx] # type: ignore
        import mimetypes
        file_name = current_file.get("name") or "Result"
        mime_type, _ = mimetypes.guess_type(file_name)
        
        # Only show the audio player if it's actually an audio/video file
        is_media = mime_type and (mime_type.startswith("audio/") or mime_type.startswith("video/"))
        is_image = mime_type and mime_type.startswith("image/")


        
        if is_media:
            st.markdown(f"Playing {idx + 1} of {len(processed_files)}: {current_file['name']}") # type: ignore
            st.audio(current_file["bytes"], format=mime_type, autoplay=True, loop=True) # type: ignore
            # Delay autoplay by 1 second so user can see results first
            if st.session_state.pop("_delay_autoplay", False):
                components.html("""
                    <script>
                        (function() {
                            var audios = window.parent.document.querySelectorAll('audio');
                            for (var i = 0; i < audios.length; i++) {
                                audios[i].pause();
                                audios[i].currentTime = 0;
                            }
                            setTimeout(function() {
                                var audios2 = window.parent.document.querySelectorAll('audio');
                                for (var j = 0; j < audios2.length; j++) {
                                    try { audios2[j].play(); } catch(e) {}
                                }
                            }, 1000);
                        })();
                    </script>
                """, height=0)
            # stats_badge_text extracted below — pass it here so it appears inline with speed controls
            _stats_for_speed = ""
            import re as _re_pre
            _stats_match = _re_pre.search(r"\*\*Statistics:\*\*\s*(.+)", display_text)
            if _stats_match:
                _stats_for_speed = _stats_match.group(1).strip()
            render_speed_controls(skill_id=selected_skill_id, stats_text=_stats_for_speed, clip_name=str(current_file.get('name', '')))
            


        elif is_image:
            st.markdown(f"Viewing Image {idx + 1} of {len(processed_files)}: {current_file['name']}") # type: ignore
            st.image(current_file["bytes"], use_container_width=True) # type: ignore
        else:
            display_name = current_file.get('original_name') # type: ignore
            if not display_name:
                # Fallback: strip .pdf from output name for document converters
                name = current_file['name'] # type: ignore
                if name.lower().endswith('.pdf'):
                    display_name = os.path.splitext(name)[0]
                else:
                    display_name = name
            # Extract stats for inline badge — API stats or native word/char count
            import re as _re_doc
            _doc_stats = ""
            _doc_stats_match = _re_doc.search(r"\*\*Statistics:\*\*\s*(.+)", display_text)
            if _doc_stats_match:
                _doc_stats = _doc_stats_match.group(1).strip()
            elif display_text.strip():
                # Native word & character count
                _word_count = len(display_text.split())
                _char_count = len(display_text)
                _doc_stats = f"{_word_count:,} words · {_char_count:,} chars"
            if _doc_stats and not is_google_sheet:
                st.markdown(f"<div style='text-align:center;'>Viewing {idx + 1} of {len(processed_files)}: {display_name}  &nbsp; <span style='background:#444;color:#ccc;padding:2px 8px;border-radius:6px;font-size:0.8em;'>{_doc_stats}</span></div>", unsafe_allow_html=True)
        # Show navigation buttons if there are multiple files
        if len(processed_files) > 1:
            label_type = "Clip" if is_media else "File"
            st.markdown("<div class='nav-button-container'>", unsafe_allow_html=True)
            col_prev, col_next = st.columns(2)
            with col_prev:
                st.markdown("<div class='nav-btn-marker' style='display:none;'></div>", unsafe_allow_html=True)
                if st.button(f"⏮ Previous {label_type}", key="prev_clip_btn", use_container_width=True):
                    # Prevent going below 0
                    set_skill_state("file_index", max(0, idx - 1))
                    set_skill_state("auto_open_result", True)
                    # Result renders as overlay — no scroll needed
                    st.rerun()
            with col_next:
                st.markdown("<div class='nav-btn-marker' style='display:none;'></div>", unsafe_allow_html=True)
                if st.button(f"Next {label_type} ⏭", key="next_clip_btn", use_container_width=True):
                    set_skill_state("file_index", min(idx + 1, len(processed_files) - 1))
                    set_skill_state("auto_open_result", True)
                    # Result renders as overlay — no scroll needed
                    st.rerun()
            st.markdown("</div>", unsafe_allow_html=True)
            
            st.components.v1.html(
                """
                <script>
                const doc = window.parent.document;
                
                // Remove existing listener to avoid duplicates on reruns
                if (doc.windowKeydownListener) {
                    doc.removeEventListener('keydown', doc.windowKeydownListener);
                }
                
                // Clear any stale focus/inline styles from previous rerun
                doc.querySelectorAll('button').forEach(b => {
                    b.blur();
                    b.classList.remove('kb-focus');
                });
                
                doc.windowKeydownListener = function(e) {
                    // Don't trigger if user is typing in an input field
                    if (e.target.tagName === 'INPUT' || e.target.tagName === 'TEXTAREA') return;
                    
                    if (e.key === 'ArrowLeft' || e.key === 'ArrowRight') {
                        const allBtns = Array.from(doc.querySelectorAll('button'));
                        // Clear highlight on ALL buttons first for a clean state
                        allBtns.forEach(b => { b.classList.remove('kb-focus'); b.blur(); });

                        // Search the button's inner text case-insensitively
                        const targetStr = e.key === 'ArrowLeft' ? 'PREVIOUS' : 'NEXT';
                        const targetBtn = allBtns.find(b => {
                            const t = b.innerText || "";
                            return t.toUpperCase().includes(targetStr);
                        });
                        
                        if (targetBtn) {
                            targetBtn.focus();
                            targetBtn.classList.add('kb-focus');
                            setTimeout(() => { targetBtn.click(); }, 300);
                        }
                    } else if (e.key === 'ArrowDown' || e.key === 'ArrowUp') {
                        // Only target download/merge buttons (with emoji icons), NOT the COPY button
                        const buttons = Array.from(doc.querySelectorAll('button')).filter(b => {
                            const text = b.innerText || "";
                            return text.includes("📥") || text.includes("📦") || text.includes("📑");
                        });
                        
                        if (buttons.length === 0) return;
                        
                        // Use the currently kb-focused button for indexing, fallback to activeElement
                        let currentIndex = buttons.indexOf(doc.querySelector('button.kb-focus'));
                        if (currentIndex === -1) currentIndex = buttons.indexOf(doc.activeElement);

                        if (e.key === 'ArrowDown') {
                            currentIndex = (currentIndex + 1) % buttons.length;
                        } else {
                            if (currentIndex <= 0) {
                                currentIndex = buttons.length - 1;
                            } else {
                                currentIndex = currentIndex - 1;
                            }
                        }
                        // Clear kb-focus and blur ALL buttons first
                        doc.querySelectorAll('button').forEach(b => { b.classList.remove('kb-focus'); b.blur(); });
                        
                        const targetBtn = buttons[currentIndex];
                        targetBtn.focus();
                        targetBtn.classList.add('kb-focus');
                        e.preventDefault();
                    } else if (e.key === 'Enter') {
                        // Click the currently kb-focused button
                        const focused = doc.querySelector('button.kb-focus');
                        if (focused) {
                            e.preventDefault();
                            e.stopImmediatePropagation();
                            const text = (focused.innerText || '').toUpperCase();
                            const link = focused.querySelector('a');
                            
                            // Top "DOWNLOAD [TYPE] NOW" button — immediate download
                            if (text.includes('DOWNLOAD') && text.includes('NOW')) {
                                if (link) {
                                    link.click();
                                } else {
                                    focused.click();
                                }
                            } else {
                                // All other buttons (ZIP, Merge) — trigger Save As dialog
                                if (link) {
                                    const origDownload = link.getAttribute('download');
                                    link.removeAttribute('download');
                                    link.click();
                                    // Restore download attribute after a tick
                                    setTimeout(() => {
                                        if (origDownload !== null) link.setAttribute('download', origDownload);
                                    }, 100);
                                } else {
                                    focused.click();
                                }
                            }
                        }
                    }
                };
                
                doc.addEventListener('keydown', doc.windowKeydownListener);
                
                // Apply specific classes to Streamlit buttons AND inject CSS into PARENT document
                if (doc.applyButtonClassesInterval) clearInterval(doc.applyButtonClassesInterval);
                doc.applyButtonClassesInterval = setInterval(() => {
                    doc.querySelectorAll('button').forEach(b => {
                        const t = (b.innerText || "").toUpperCase();
                        if (t.includes('DOWNLOAD ALL')) {
                            b.classList.add('merge-btn-cyan');
                        }
                        if (t.includes('PREVIOUS') || t.includes('NEXT')) {
                            b.classList.add('nav-btn-15px');
                            // On mouseover, clear keyboard focus class from ALL nav buttons
                            if (!b._hoverBound) {
                                b._hoverBound = true;
                                b.addEventListener('mouseenter', () => {
                                    doc.querySelectorAll('button.nav-btn-15px').forEach(nb => {
                                        nb.classList.remove('kb-focus');
                                    });
                                });
                            }
                        }
                        // Download/merge buttons — clear keyboard focus class on hover
                        if (t.includes('📥') || t.includes('📦') || t.includes('📑')) {
                            if (!b._dlHoverBound) {
                                b._dlHoverBound = true;
                                b.addEventListener('mouseenter', () => {
                                    doc.querySelectorAll('button').forEach(ob => {
                                        const ot = ob.innerText || '';
                                        if (ot.includes('📥') || ot.includes('📦') || ot.includes('📑')) {
                                            ob.classList.remove('kb-focus');
                                        }
                                    });
                                });
                            }
                        }
                    });
                }, 100);
                
                // Inject CSS into the PARENT document (not the iframe) so it actually applies
                if (!doc.getElementById('ag-custom-btn-styles')) {
                    const style = doc.createElement('style');
                    style.id = 'ag-custom-btn-styles';
                    style.textContent = `
                        /* White outline on hover only - focus-visible for keyboard only */
                        .stApp button:hover {
                            border: 2px solid #ffffff !important;
                            box-shadow: 0 0 0 4px rgba(255, 255, 255, 0.85) !important;
                            outline: none !important;
                        }
                        .stApp button:focus {
                            outline: none !important;
                            border-color: transparent !important;
                            box-shadow: none !important;
                        }
                        /* Keyboard focus class - overrides :focus suppression */
                        .stApp button.kb-focus {
                            border: 2px solid #ffffff !important;
                            box-shadow: 0 0 0 4px rgba(255, 255, 255, 0.85) !important;
                        }
                        /* Specific Cyan override for Merge buttons - all states */
                        .stApp button.merge-btn-cyan,
                        .stApp button.merge-btn-cyan:hover,
                        .stApp button.merge-btn-cyan:active,
                        .stApp button.merge-btn-cyan:focus {
                            background-color: #00FFFF !important;
                            color: black !important;
                            border-color: #00cccc !important;
                        }
                        .stApp button.merge-btn-cyan:hover {
                            background-color: #00e5e5 !important;
                            border: 2px solid #ffffff !important;
                            box-shadow: 0 0 0 4px rgba(255, 255, 255, 0.85) !important;
                        }
                        .stApp button.merge-btn-cyan p,
                        .stApp button.merge-btn-cyan div,
                        .stApp button.merge-btn-cyan:hover p,
                        .stApp button.merge-btn-cyan:hover div {
                            color: black !important;
                        }
                    `;
                    doc.head.appendChild(style);
                }
                </script>
                """,
                height=0
            )

        # --- Quick Select List & Upload (Playlist) ---
        # Hide for Google Sheet skill as it's redundant (always just one "Result")
        if not is_google_sheet:
            # Make expander text 40% larger and position text above arrow
            st.markdown("<h4 style='text-align:left; margin-top:2em; margin-bottom:1em;'>SHOW MY DOCS</h4>", unsafe_allow_html=True)
            st.markdown("""<style>
                .stApp:has(#ag-result-marker) div[data-testid="stExpander"] summary {
                    flex-direction: column-reverse !important;
                    align-items: center !important;
                    gap: 0px !important;
                }
                .stApp:has(#ag-result-marker) div[data-testid="stExpander"] summary span {
                    font-size: 1.4em !important;
                }
            </style>""", unsafe_allow_html=True)
            with st.expander("", expanded=False):
                if is_media:
                    st.markdown("<h3 class='centered-header'>MY CLIPS</h3>", unsafe_allow_html=True)
                else:
                    pass
                
                if get_skill_state("popup_batch_success"):
                    st.markdown("<div class='success-message-popup'>✅ SUCCESSFULLY ADDED!</div>", unsafe_allow_html=True)
                    set_skill_state("popup_batch_success", False)
                
                st.markdown("<div class='quick-select-btns'>", unsafe_allow_html=True)
                for i, clip in enumerate(processed_files):
                    is_active = (i == get_skill_state("file_index", 0))
                    if is_active:
                        st.markdown(f"<div class='active-clip'>🟢 {clip['name']}</div>", unsafe_allow_html=True)
                    else:
                        with st.container():
                            st.markdown("<div class='playlist-clip-marker'></div>", unsafe_allow_html=True)
                            if st.button(f"⚪️ {clip['name']}", key=f"clip_btn_{i}", use_container_width=True):
                                set_skill_state("file_index", i)
                                set_skill_state("auto_open_result", True)
                                st.rerun()
                st.markdown("</div>", unsafe_allow_html=True)

    # Render the transcription box BELOW the navigation buttons
    import html
    import re as _re

    # Extract and strip the Statistics line from display_text so it shows ONLY in the badge
    stats_badge_text = ""
    stats_match = _re.search(r"\*\*Statistics:\*\*\s*(.+)", display_text)
    if stats_match:
        stats_badge_text = stats_match.group(1).strip()
        display_text = _re.sub(r"\n\n\*\*Statistics:\*\*\s*.+", "", display_text).strip()
    
    # --- COPY button ABOVE preview ---
    st.markdown("<div style='margin-top:1.5em;'></div>", unsafe_allow_html=True)
    if is_google_sheet:
        import re
        sheet_title = "Google Sheet"
        sheet_url = "#"
        
        # Strip ANSI escape codes that might be polluting the terminal stream
        clean_text = _re.sub(r'\x1b\[[0-9;]*[mG]', '', display_text)
        
        # Extract title from the script's output "Creating Google Sheet: '...'"
        title_match = re.search(r"Creating Google Sheet:\s*['\"]?(.*?)['\"]?\s*(?:\.\.\.|\n|$)", clean_text)
        if title_match:
            sheet_title = title_match.group(1).strip()
            
        # Extract the URL from the script's output
        url_match = re.search(r"https://docs\.google\.com/spreadsheets/d/[a-zA-Z0-9_-]+[^\s\"']*", clean_text)
        if url_match:
            sheet_url = url_match.group(0).strip()
            
        st.markdown(
            f"""
            <div style="display: flex; justify-content: center; margin-top: 10px; margin-bottom: 20px;">
                <a href="{sheet_url}" target="_blank" style="
                    background-color: #1e1e1e;
                    color: #00FFCC;
                    border: 2px solid #00FFCC;
                    border-radius: 6px;
                    padding: 8px 24px;
                    font-size: 1.2rem;
                    font-family: sans-serif;
                    font-weight: bold;
                    text-align: center;
                    text-decoration: none;
                    box-shadow: 0 4px 6px rgba(0, 0, 0, 0.3);
                    transition: background-color 0.2s, transform 0.1s;
                " onmouseover="this.style.backgroundColor='#003322'; this.style.transform='scale(1.02)';" onmouseout="this.style.backgroundColor='#1e1e1e'; this.style.transform='scale(1)';">
                    📊 {sheet_title}
                </a>
            </div>
            """, unsafe_allow_html=True
        )
    else:
        import json
        js_escaped_text = json.dumps(display_text)
        st.components.v1.html(
            f"""
            <div style="display: flex; justify-content: center; margin-top: 5px; margin-bottom: 15px;">
                <button id="copy-btn" style="
                    background-color: #ffe700;
                    color: #000000;
                    border: 1px solid #ffe700;
                    border-radius: 4px;
                    padding: 0px 16px;
                    height: 38px;
                    font-size: 1rem;
                    cursor: pointer;
                    font-family: sans-serif;
                    font-weight: bold;
                    transition: all 0.2s;
                    min-width: 100px;
                    text-transform: uppercase;
                ">COPY</button>
            </div>
            <script>
            const btn = document.getElementById('copy-btn');
            btn.onclick = function() {{
                navigator.clipboard.writeText({js_escaped_text}).then(() => {{
                    btn.innerText = "\\u2713 COPIED!";
                    btn.style.backgroundColor = "#8cd775";
                    btn.style.borderColor = "#8cd775";
                    setTimeout(() => {{
                        btn.innerText = "COPY";
                        btn.style.backgroundColor = "#ffe700";
                        btn.style.borderColor = "#ffe700";
                    }}, 2000);
                }});
            }};
            btn.onmouseover = function() {{ this.style.backgroundColor = '#ffd600'; this.style.borderColor = '#ffd600'; }};
            btn.onmouseout = function() {{
                if (this.innerText === "COPY") {{
                    this.style.backgroundColor = '#ffe700';
                    this.style.borderColor = '#ffe700';
                }}
            }};
            </script>
            """,
            height=70,
        )

    # --- Preview box BELOW copy button ---
    if is_media or is_image:
        # Check if word tracker alignment data is available for this file
        _wt_alignment = None
        if processed_files:
            _wt_current = processed_files[min(idx, len(processed_files)-1)]
            _wt_alignment = _wt_current.get("alignment")
        
        if _wt_alignment and isinstance(_wt_alignment, dict) and _wt_alignment.get("words"):
            # Render synchronized word tracker instead of static text
            render_word_tracker(_wt_alignment, clip_name=str(current_file.get('name', '') if processed_files else ''))
        elif display_text:
            # Use estimated word tracker for media files with transcript but no alignment data
            # This provides Follow Along parity between Speech2Text and Text2Speech
            _est_words = display_text.split()
            if _est_words and is_media:
                _est_alignment = {"words": _est_words, "start_times": [], "end_times": []}
                render_word_tracker(_est_alignment, clip_name=str(current_file.get('name', '') if processed_files else ''), estimated=True)
            else:
                safe_text = html.escape(display_text)
                st.markdown(
                    f"<div class='transcript-box'>{safe_text.replace(chr(10), '<br>')}</div>",
                    unsafe_allow_html=True
                )
    else:
        # Use the same transcript-box styling for document previews
        if display_text and not is_google_sheet:
            safe_text = html.escape(display_text)
            st.markdown(
                f"<div class='transcript-box'>{safe_text.replace(chr(10), '<br>')}</div>",
                unsafe_allow_html=True
            )
    


    # --- Dynamic Direct Download Support (Top level, synced to active file) ---
    if processed_files:
        current_item = processed_files[idx]
        res_bytes = current_item.get("result_bytes")
        res_name = current_item.get("result_name")
        
        if res_bytes and res_name:
            st.markdown("<br>", unsafe_allow_html=True)
            
            # Force correct extension if this is the PDF Converter
            if selected_skill["basename"] == "Convtr-PlainTxt2PDF" and not res_name.lower().endswith(".pdf"):
                res_name = os.path.splitext(res_name)[0] + ".pdf"
            
            # Use mimetypes for clean extension labeling
            import mimetypes
            dl_mime, _ = mimetypes.guess_type(res_name)
            ext_label = os.path.splitext(res_name)[1][1:].upper()
            
            # Determine dynamic button label
            dl_btn_label = "DOWNLOAD PDF" if selected_skill["basename"] == "Convtr-PlainTxt2PDF" else f"DOWNLOAD {ext_label}"
            
            # Specialized player for audio results (like TTS)
            if dl_mime and dl_mime.startswith("audio/"):
                st.audio(res_bytes, format=dl_mime, autoplay=True) # type: ignore
                render_speed_controls(skill_id=selected_skill_id, clip_name=str(res_name or ''))
                st.markdown("<br>", unsafe_allow_html=True)
                
            # We strictly bypass both Streamlit's native button bug AND Chrome's iframe blob sandbox bug
            # by dynamically saving the generated byte sequence as a PHYSICAL static file on the server.
            # This allows us to provide Chrome with a fully standard HTTP GET request to a URL ending internally
            # with the exact `.pdf` filename, completely shielding it from Chrome's payload security heuristics.
            
            static_dir = os.path.join(os.path.dirname(__file__), "static", "downloads")
            os.makedirs(static_dir, exist_ok=True)
            
            # Ensure safe filename
            import urllib.parse
            safe_res_name = "".join([c for c in res_name if c.isalpha() or c.isdigit() or c in (' ', '.', '-', '_')]).rstrip()
            file_path = os.path.join(static_dir, safe_res_name)
            
            with open(file_path, "wb") as f:
                f.write(res_bytes)
            
            # The static path mapped by config.toml [server] enableStaticServing
            static_url = f"app/static/downloads/{urllib.parse.quote(safe_res_name)}"
            
            st.markdown(f"""
                <div class="custom-html-download-btn-container" style="margin-bottom: 2rem; width: 100%;">
                    <a href="{static_url}" download="{safe_res_name}" data-customdl="true" style="background-color: #8cd775; color: #000000; border: 1px solid #8cd775; border-radius: 4px; padding: 0.5rem 1rem; height: auto; min-height: 2.5rem; font-size: 1rem; cursor: pointer; font-weight: bold; transition: all 0.2s; text-transform: uppercase; box-sizing: border-box; display: inline-flex; align-items: center; justify-content: center; text-decoration: none; font-family: 'Source Sans Pro', sans-serif;">
                        📥 {dl_btn_label}
                    </a>
                </div>
            """, unsafe_allow_html=True)
            
            # Use JS to manipulate the parent DOM container natively per streamlit-styling.md workflow!
            components.html("""
            <script>
            (function styleDownloadBtnContainer() {
                const links = window.parent.document.querySelectorAll('a[data-customdl="true"]');
                links.forEach(link => {
                    const stMarkdown = link.closest('div[data-testid="stMarkdownContainer"]');
                    if (stMarkdown) {
                        stMarkdown.style.setProperty("display", "flex", "important");
                        stMarkdown.style.setProperty("justify-content", "center", "important");
                        stMarkdown.style.setProperty("width", "100%", "important");
                        
                        const elementContainer = stMarkdown.closest('div.element-container');
                        if (elementContainer) {
                            elementContainer.style.setProperty("display", "flex", "important");
                            elementContainer.style.setProperty("justify-content", "center", "important");
                            elementContainer.style.setProperty("width", "100%", "important");
                        }
                    }
                });
            })();
            </script>
            """, height=0)
            
            # --- TOP BATCH ACTIONS (For PDF Converter) ---
            if selected_skill["basename"] == "Convtr-PlainTxt2PDF" and len(processed_files) > 1: # type: ignore
                st.markdown("<div style='margin-bottom: 10px;'></div>", unsafe_allow_html=True)
                
                # Use same column constraints as "DOWNLOAD PDF" to match button size
                col_b1, col_b2, col_b3 = st.columns([2, 3, 2])
                with col_b2:
                    # 1. Top Buttons: ZIP All (Cyan overridden in JS/Tertiary base)
                    zip_bytes_clean = generate_zip_of_all_transcripts(processed_files, "PDF (.pdf)", include_sources=False)
                    if zip_bytes_clean:
                        st.download_button(label=f"📦 DOWNLOAD ALL (CLEAN)", data=zip_bytes_clean, file_name="All_Files_Clean.zip", mime="application/zip", use_container_width=True, type="tertiary", key="popup_dl_all_zip_clean_top")
                    
                    zip_bytes_sourced = generate_zip_of_all_transcripts(processed_files, "PDF (.pdf)", include_sources=True)
                    if zip_bytes_sourced:
                        st.download_button(label=f"📦 DOWNLOAD ALL (INCLUDE SOURCES)", data=zip_bytes_sourced, file_name="All_Files_Sourced.zip", mime="application/zip", use_container_width=True, type="tertiary", key="popup_dl_all_zip_sourced_top")

                    # 2. Bottom Buttons: Merge All (Orange/Tertiary)
                    merged_parts_clean = []
                    merged_parts_sourced = []
                    for f in processed_files:
                        content = f.get("content_preview") or ""
                        if not content.strip():
                            transcript = f.get("transcript", "")
                            if not any(x in transcript for x in ["Successfully", "Generated:", "✅"]):
                                content = transcript
                        if content.strip():
                            merged_parts_clean.append(content)
                            merged_parts_clean.append("\n")
                            merged_parts_sourced.append(f"Sourced from: {f['name']}") # Matches orange styling in script
                            merged_parts_sourced.append("\n\n")
                            merged_parts_sourced.append(content)
                            merged_parts_sourced.append("\n")
                    
                    merged_bytes_clean = generate_pdf_from_text("\n".join(merged_parts_clean))
                    merged_bytes_sourced = generate_pdf_from_text("\n".join(merged_parts_sourced))
                    if merged_bytes_clean:
                        st.download_button(label=f"📑 MERGE ALL (CLEAN)", data=merged_bytes_clean, file_name="Merged_Document_Clean.pdf", mime="application/pdf", use_container_width=True, type="tertiary", key="popup_merge_pdf_clean_top") # type: ignore
                    if merged_bytes_sourced:
                        st.download_button(label=f"📑 MERGE ALL (INCLUDE SOURCES)", data=merged_bytes_sourced, file_name="Merged_Document_Sourced.pdf", mime="application/pdf", use_container_width=True, type="tertiary", key="popup_merge_pdf_sourced_top") # type: ignore


    # --- ADD MORE Files Button (inside expander) ---
    st.markdown("<div style='margin-top: 20px;'></div>", unsafe_allow_html=True)
    # Determine allowed types based on current skill
    _addmore_types = None
    _sel_skill_raw = st.session_state.get("selected_skill_id", "")
    _is_audio_addmore = any(n in str(_sel_skill_raw) for n in ["Speech2Text", "KIE-ElevenLabs-Speech2Text"])
    _is_tts_addmore = any(n in str(_sel_skill_raw) for n in ["Text2Speech", "KIE-ElevenLabs-Text2Speech"])
    if _is_audio_addmore:
        _addmore_types = ["mp3", "wav", "m4a", "aac", "ogg", "flac", "webm", "aiff", "aif", "wma", "oga", "opus", "3gp", "mp4", "mov", "avi", "mkv"]
    elif _is_tts_addmore:
        _addmore_types = ["txt", "md", "rtf", "doc", "docx", "csv", "json", "py", "sh", "yaml", "yml", "ini"]
    else:
        _addmore_types = ["txt", "md", "docx", "doc", "csv", "json", "rtf", "py", "sh", "yaml", "yml"]

    # --- ADD MORE FILES: Native Streamlit button styled with exact CSS classes injected via JS ---
    _addmore_counter = st.session_state.get("_addmore_key_counter", 0)
    _addmore_open_key = "addmore_panel_open"

    st.markdown("""
        <style>
            /* Explicit classes added by the Javascript below */
            .custom-addmore-btn {
                background-color: transparent !important;
                border: none !important;
                color: #ffd700 !important;
                font-size: 1rem !important;
                font-family: 'Source Sans Pro', sans-serif !important;
                font-weight: 700 !important;
                padding: 0.5em 1em !important;
                transition: all 0.2s !important;
                box-shadow: none !important;
                min-height: 0 !important;
                height: auto !important;
                outline: none !important;
            }
            .custom-addmore-btn:hover {
                opacity: 0.8 !important;
                background-color: transparent !important;
                color: #ffd700 !important;
                border: none !important;
                outline: none !important;
            }
            .custom-addmore-btn p {
                color: inherit !important;
                font-weight: 700 !important;
                margin: 0 !important;
                text-align: center !important;
                width: 100% !important;
            }
            .custom-addmore-btn div {
                justify-content: center !important;
                display: flex !important;
                width: 100% !important;
            }
        </style>
    """, unsafe_allow_html=True)
    
    st.markdown("<br>", unsafe_allow_html=True)  # spacing
    # use_container_width=True forces it 100% wide, allowing the JS text styling to perfectly center it visually!
    if st.button("➕ ADD MORE FILES", key="addmore_toggle_btn", use_container_width=True):
        st.session_state[_addmore_open_key] = not st.session_state.get(_addmore_open_key, False)
        st.rerun()


    # Show file uploader when panel is open
    _addmore_files = None
    if st.session_state.get(_addmore_open_key, False):
        _addmore_files = st.file_uploader(
            "Upload additional files",
            accept_multiple_files=True,
            label_visibility="collapsed",
            type=_addmore_types,
            key=f"addmore_uploader_{_addmore_counter}",
        )

    if _addmore_files:
        # --- Duplicate filtering: reject files already processed or already in the upload queue ---
        _processed_set = get_skill_state("processed_files", set())
        _processed_set = _processed_set if isinstance(_processed_set, set) else set()
        # Also check against whatever is currently in the main uploader
        _current_uploaded = st.session_state.get(f"file_uploader_{st.session_state.get('selected_skill_id','')}", None)
        _existing_ids = set()
        if _current_uploaded:
            for _uf in (_current_uploaded if isinstance(_current_uploaded, list) else [_current_uploaded]):
                _existing_ids.add(_uf.name + str(_uf.size))
        # Also include already-processed file IDs
        _existing_ids.update(_processed_set)  # type: ignore

        _new_files = []
        _skipped = []
        for _af in _addmore_files:
            file_id = _af.name + str(_af.size)
            if file_id in _existing_ids:
                _skipped.append(_af.name)
            else:
                _new_files.append(_af)

        if _skipped:
            trigger_duplicate_error()
            # Clear the newly added files to prevent infinite rerun loop of triggering error
            _addmore_counter += 1
            st.session_state["_addmore_key_counter"] = _addmore_counter
            st.rerun()
            
        elif _new_files:
            # We have valid new files -> trigger main processing and clean up UI
            set_skill_state("addmore_pending_files", _new_files)
            # Hide uploader and clear its state to reset for next time
            st.session_state[_addmore_open_key] = False
            _addmore_counter += 1
            st.session_state["_addmore_key_counter"] = _addmore_counter
            
            # Since we are triggering a new backend run, ensure any auto-open triggers are removed 
            # so the success overlay doesn't randomly pop open immediately.
            set_skill_state("auto_open_result", False)
            st.rerun()

    # --- CLOSE BUTTON (Final UX styled, properly centered via explicit JS classes) ---
    st.markdown("""
        <style>
            /* Explicit classes added by the Javascript below */
            .custom-clear-btn {
                background-color: transparent !important;
                color: #aaaaaa !important;
                border: 1px solid #aaaaaa !important;
                border-radius: 4px !important;
                padding: 0px !important;
                height: 38px !important;
                width: 100px !important;
                min-width: 100px !important;
                max-width: 100px !important;
                font-size: 0.9rem !important;
                font-family: 'Source Sans Pro', sans-serif !important;
                transition: all 0.2s !important;
                box-shadow: none !important;
                font-weight: normal !important;
                outline: none !important;
                text-align: center !important;
                white-space: nowrap !important;
            }
            .custom-clear-btn:hover {
                color: #ffffff !important;
                border-color: #ffffff !important;
                background-color: transparent !important;
                outline: none !important;
            }
            .custom-clear-btn p {
                color: inherit !important;
                margin: 0 !important;
                text-align: center !important;
                white-space: nowrap !important;
            }
            .custom-clear-btn div {
                justify-content: center !important;
                display: flex !important;
            }
        </style>
    """, unsafe_allow_html=True)
    
    # Use a specific height div to explicitly enforce the spacing above. Increased to move down 1 line.
    st.markdown("<div style='height: 3.5em;'></div>", unsafe_allow_html=True)
    
    # We use use_container_width=False so it never word-wraps, and perfectly right-justify
    # it safely using text-align in the javascript block below, avoiding flex overlaps.
    if st.button("CLEAR ALL", key="close_popup_final_ux", use_container_width=False):
        set_skill_state("last_output", None)
        set_skill_state("auto_open_result", None)
        set_skill_state("direct_download_file", None)
        st.rerun()

    # --- JAVASCRIPT POST-RENDER STYLING ALGORITHM ---
    # According to the Streamlit styling workflow, the only bulletproof way to target elements 
    # without failing due to nested CSS block specificity or testid mutations, is matching text content via JS:
    components.html("""
    <script>
    (function styleButtons() {
        const doc = window.parent.document;
        let elements = Array.from(doc.querySelectorAll('.stButton p'));
        let foundAdd = false;
        let foundClear = false;
        
        for (const p of elements) {
            if (p.textContent.includes('ADD MORE FILES')) {
                const btn = p.closest('button');
                if (btn) {
                    btn.classList.add('custom-addmore-btn');
                    // Ensure the button parent container is set to full width block display so text aligns center
                    const wrapper = btn.closest('div[data-testid="stElementContainer"]');
                    if (wrapper) wrapper.style.cssText = 'width: 100% !important; display: block !important;';
                    foundAdd = true;
                }
            }
            if (p.textContent.includes('CLEAR ALL')) {
                const btn = p.closest('button');
                if (btn) {
                    btn.classList.add('custom-clear-btn');
                    // Force Streamlit's native button wrapper to stop auto-expanding
                    const stButton = btn.closest('.stButton');
                    if (stButton) {
                        stButton.style.setProperty("width", "100px", "important");
                        stButton.style.setProperty("min-width", "100px", "important");
                        stButton.style.setProperty("max-width", "100px", "important");
                        stButton.style.setProperty("flex", "0 0 100px", "important");
                    }
                    // According to the styling workflow manual, use robust Flexbox property assignments
                    const stElement = btn.closest('div[data-testid="stElementContainer"]');
                    if (stElement) {
                        stElement.style.setProperty("display", "flex", "important");
                        stElement.style.setProperty("justify-content", "flex-end", "important");
                        stElement.style.setProperty("width", "100%", "important");
                    }
                    const elContainer = btn.closest('div.element-container');
                    if (elContainer) {
                        elContainer.style.setProperty("display", "flex", "important");
                        elContainer.style.setProperty("justify-content", "flex-end", "important");
                        elContainer.style.setProperty("width", "100%", "important");
                    }
                    foundClear = true;
                }
            }
        }
        
        if (!foundAdd || !foundClear) {
            setTimeout(styleButtons, 200);
        }
    })();
    </script>
    """, height=0)




# --- RENDER RESULT AS CENTERED POPUP OVERLAY ---
last_output = get_skill_state("last_output")

# Always inject the overlay CSS into parent document (CSS :has() auto-activates on marker presence)
# This MUST run before the result renders so the style is ready when the marker appears
components.html("""
    <script>
        (function() {
            var doc = window.parent.document;
            // Clean up any leftover curtains from previous approach
            var curtain = doc.getElementById('ag-scroll-curtain');
            if (curtain) curtain.remove();
            
            // Inject (or refresh) the overlay CSS — always replace to pick up code changes
            var existingStyle = doc.getElementById('ag-result-overlay-css');
            if (existingStyle) existingStyle.remove();
            {
                var style = doc.createElement('style');
                style.id = 'ag-result-overlay-css';
                style.textContent = `
                    /* When #ag-result-marker exists, make the app viewport a fixed overlay */
                    .stApp:has(#ag-result-marker) > div[data-testid="stAppViewContainer"] {
                        position: fixed !important;
                        top: 0 !important;
                        left: 0 !important;
                        width: 100vw !important;
                        height: 100vh !important;
                        z-index: 999970 !important;
                        background: rgba(0,0,0,0.92) !important;
                        overflow-y: auto !important;
                        overflow-x: hidden !important;
                        padding: 3vh 8vw !important;
                    }
                    /* Hide the header/toolbar behind the overlay */
                    .stApp:has(#ag-result-marker) > header {
                        z-index: 0 !important;
                    }
                    /* HIDE all page content that is NOT the result popup */
                    /* Target top-level children in the main vertical block */
                    .stApp:has(#ag-result-marker) div[data-testid="stVerticalBlock"] > div {
                        display: none !important;
                    }
                    /* But SHOW the container that holds the result marker (and all its children) */
                    .stApp:has(#ag-result-marker) div[data-testid="stVerticalBlock"] > div:has(#ag-result-marker) {
                        display: block !important;
                    }
                    /* Un-hide all descendants inside the result popup */
                    .stApp:has(#ag-result-marker) div[data-testid="stVerticalBlock"] > div:has(#ag-result-marker) *:not(style):not(script) {
                        display: revert !important;
                    }
                    /* Keep style/script tags invisible */
                    .stApp:has(#ag-result-marker) div:has(#ag-result-marker) style,
                    .stApp:has(#ag-result-marker) div:has(#ag-result-marker) script {
                        display: none !important;
                    }

                    /* Remove visibility of the hidden streamlit buttons used for routing */
                    .stApp:has(#ag-result-marker) div[data-testid="element-container"]:has(#hide-clear-btn-marker),
                    .stApp:has(#ag-result-marker) div[data-testid="element-container"]:has(#hide-clear-btn-marker) + div[data-testid="element-container"],
                    .stApp:has(#ag-result-marker) div[data-testid="element-container"]:has(#hide-addmore-btn-marker),
                    .stApp:has(#ag-result-marker) div[data-testid="element-container"]:has(#hide-addmore-btn-marker) + div[data-testid="element-container"] {
                        position: absolute !important;
                        opacity: 0 !important;
                        pointer-events: none !important;
                        height: 1px !important;
                        width: 1px !important;
                        overflow: hidden !important;
                        margin: 0 !important;
                        padding: 0 !important;
                    }

                    /* Hide Streamlit heading anchor link icon in the result popup */
                    .stApp:has(#ag-result-marker) div:has(#ag-result-marker) h1 a,
                    .stApp:has(#ag-result-marker) div:has(#ag-result-marker) [data-testid="StyledLinkIconContainer"] {
                        display: none !important;
                    }
                    
                    /* Fix file uploader appearance in popup (replacing removed JS) */
                    .stApp:has(#ag-result-marker) div:has(#ag-result-marker) [data-testid="stFileUploader"] [data-testid="stWidgetLabel"] {
                        display: none !important;
                    }
                    .stApp:has(#ag-result-marker) div:has(#ag-result-marker) section[data-testid="stFileUploader"] {
                        margin-bottom: 1em !important;
                    }
                `;
                doc.head.appendChild(style);
            }
            
            // Periodically hunt down and destroy the visibility of the UX_HIDDEN buttons
            // because CSS :has() + adjacent sibling combinators are too brittle against
            // Streamlit's dynamic DOM node generation and `display: revert !important` rule.
            function hideUxButtons() {
                var btns = doc.querySelectorAll('button');
                for (var i = 0; i < btns.length; i++) {
                    if (btns[i].textContent.indexOf('UX_HIDDEN') !== -1) {
                        // Hide the button itself
                        btns[i].style.setProperty('display', 'none', 'important');
                        // Hide its immediate wrapper
                        var wrapper = btns[i].closest('[data-testid="element-container"]');
                        if (wrapper) wrapper.style.setProperty('display', 'none', 'important');
                    }
                }
            }
            hideUxButtons();
            setInterval(hideUxButtons, 500); // Keep them hidden even if Streamlit re-renders

        })();
    </script>
""", height=0)

if last_output:
    # Render the result content in a container with a marker
    with st.container():
        # Hidden marker — CSS :has() selector auto-activates the overlay when this exists
        st.markdown("<div id='ag-result-marker' style='display:none;'></div>", unsafe_allow_html=True)
        show_result_popup(last_output)
