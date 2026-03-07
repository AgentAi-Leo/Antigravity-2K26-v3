from docx import Document
document = Document()
document.add_paragraph('First item in bullet list', style='List Bullet')
document.add_paragraph('Second item in bullet list', style='List Bullet')
document.add_paragraph('First item in numbered list', style='List Number')
document.add_paragraph('    Indented item', style='Normal')
document.save('test.docx')
